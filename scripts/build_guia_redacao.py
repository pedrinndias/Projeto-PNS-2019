# -*- coding: utf-8 -*-
"""Gera o guia de redação do artigo científico.

NÃO é o artigo. É um andaime: para cada seção do modelo de artigo do Prof. Zárate,
o guia mostra (a) o que o gabarito exige, (b) qual material já produzido nos
notebooks alimenta aquela seção, (c) perguntas norteadoras para o autor pensar e
descobrir o que escrever, e (d) o que ainda falta.

Uso (rodar da raiz do projeto):
    python scripts/build_guia_redacao.py

Gera: Documentos_organizacao/Guia_Redacao_Artigo_PNS2019.docx
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT_PATH = ROOT / "Documentos_organizacao" / "Guia_Redacao_Artigo_PNS2019.docx"

AZUL_ESCURO = RGBColor(0x1F, 0x37, 0x64)
AZUL_MEDIO  = RGBColor(0x2E, 0x75, 0xB6)
CINZA       = RGBColor(0x59, 0x59, 0x59)
VERDE       = RGBColor(0x1E, 0x7A, 0x34)
LARANJA     = RGBColor(0xB7, 0x5E, 0x09)
VERMELHO    = RGBColor(0xA6, 0x1B, 0x1B)


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────
def setup_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    for lvl, size, color in [(1, 17, AZUL_ESCURO), (2, 13.5, AZUL_MEDIO), (3, 12, AZUL_MEDIO)]:
        st = doc.styles[f"Heading {lvl}"]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color


def h1(doc, txt): return doc.add_heading(txt, level=1)
def h2(doc, txt): return doc.add_heading(txt, level=2)
def h3(doc, txt): return doc.add_heading(txt, level=3)


def para(doc, txt, italic=False, size=11, bold=False, justify=True, color=None):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(txt)
    r.italic = italic
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p


def bullet(doc, txt, bold_prefix=None, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    if bold_prefix:
        rb = p.add_run(f"{bold_prefix} ")
        rb.bold = True
    p.add_run(txt)
    return p


def numbered(doc, txt, bold_prefix=None):
    p = doc.add_paragraph(style="List Number")
    if bold_prefix:
        rb = p.add_run(f"{bold_prefix} ")
        rb.bold = True
    p.add_run(txt)
    return p


def question(doc, txt):
    """Pergunta norteadora — formatação distinta."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(4)
    rb = p.add_run("?  ")
    rb.bold = True
    rb.font.color.rgb = AZUL_MEDIO
    r = p.add_run(txt)
    r.font.size = Pt(10.5)
    return p


def _shade(cell, hex_color):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): hex_color})
    tcpr.append(shd)


def callout(doc, titulo, linhas, cor="FFF3CD"):
    """Caixa de destaque com título e uma ou mais linhas."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    _shade(cell, cor)
    cell.text = ""
    p1 = cell.paragraphs[0]
    r1 = p1.add_run(titulo)
    r1.bold = True
    r1.font.size = Pt(10.5)
    if isinstance(linhas, str):
        linhas = [linhas]
    for ln in linhas:
        p = cell.add_paragraph()
        r = p.add_run(ln)
        r.font.size = Pt(10)
    doc.add_paragraph()


def status_table(doc, rows, col_widths=(1.2, 7.8, 6.5)):
    """Tabela de 3 colunas: status | item | observação."""
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, txt in enumerate(["Status", "Item", "Onde está / o que fazer"]):
        hdr[i].text = txt
        for par in hdr[i].paragraphs:
            for run in par.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade(hdr[i], "2E75B6")
    cores = {"PRONTO": "D5EED5", "PARCIAL": "FFF0CC", "PENDENTE": "F5D5D5"}
    for st, item, obs in rows:
        cells = table.add_row().cells
        cells[0].text = st
        cells[1].text = item
        cells[2].text = obs
        _shade(cells[0], cores.get(st, "FFFFFF"))
        for i, c in enumerate(cells):
            for par in c.paragraphs:
                for run in par.runs:
                    run.font.size = Pt(8.5)
                    if i == 0:
                        run.bold = True
    for row in table.rows:
        for i, w in enumerate(col_widths):
            row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def map_table(doc, rows, headers, col_widths):
    """Tabela genérica de mapeamento."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, txt in enumerate(headers):
        hdr[i].text = txt
        for par in hdr[i].paragraphs:
            for run in par.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade(hdr[i], "1F3764")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for par in cells[i].paragraphs:
                for run in par.runs:
                    run.font.size = Pt(8.5)
    for row in table.rows:
        for i, w in enumerate(col_widths):
            row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def pagebreak(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# ─────────────────────────────────────────────────────────────────────────────
# Construção
# ─────────────────────────────────────────────────────────────────────────────
def build():
    doc = Document()
    setup_styles(doc)
    for s in doc.sections:
        s.top_margin = Cm(2.2); s.bottom_margin = Cm(2.2)
        s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

    # ===== CAPA =====
    for _ in range(4):
        doc.add_paragraph()
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Guia de Redação do Artigo Científico")
    r.bold = True; r.font.size = Pt(26); r.font.color.rgb = AZUL_ESCURO
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Padrões Nutricionais e Artrite em Idosos Brasileiros — PNS 2019")
    rs.italic = True; rs.font.size = Pt(14); rs.font.color.rgb = CINZA
    doc.add_paragraph()
    ln = doc.add_paragraph(); ln.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ln.add_run("_" * 50).font.color.rgb = AZUL_MEDIO
    for _ in range(2):
        doc.add_paragraph()
    msg = doc.add_paragraph(); msg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rm = msg.add_run("Este documento NÃO é o artigo.\nÉ o andaime que organiza tudo o que já produzimos "
                     "e faz as perguntas certas para você\ndescobrir, sozinho, o que escrever em cada seção.")
    rm.font.size = Pt(12); rm.font.color.rgb = AZUL_ESCURO; rm.bold = True
    for _ in range(6):
        doc.add_paragraph()
    for rotulo, valor in [
        ("Autor", "Pedro Dias Soares"),
        ("Orientador", "Prof. Dr. Luis Enrique Zárate — PUC Minas"),
        ("Modelo-base", "padroes_nutricao_artrite_idosos.docx (template do Prof. Zárate)"),
        ("Plano de ação", "Plano_Artigo_Mineracao_Pedro_Dias_Soares.docx"),
        ("Material consolidado", "Resultados_Consolidados_PNS2019_Artrite.docx"),
    ]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = p.add_run(f"{rotulo}:  "); rr.bold = True; rr.font.size = Pt(10)
        p.add_run(valor).font.size = Pt(10)
    pagebreak(doc)

    # ===== COMO USAR =====
    h1(doc, "Como usar este guia")
    para(doc,
         "O modelo de artigo do Prof. Zárate tem caixas com placeholders [X] esperando que você "
         "preencha. O problema não é preencher — é saber O QUE escrever. Este guia resolve isso. "
         "Para cada seção do artigo, você encontra quatro blocos:")
    bullet(doc, "explica o que aquela seção precisa conter, segundo o checklist de 29 critérios.",
           bold_prefix="O que o gabarito exige —")
    bullet(doc, "lista exatamente quais tabelas, figuras, números e discussões que já produzimos "
                "alimentam aquela seção, e onde encontrá-los.", bold_prefix="Material que já temos —")
    bullet(doc, "uma lista de perguntas. Responder cada uma, com suas palavras, É o rascunho da seção. "
                "Não pule esta parte — é o coração do guia.", bold_prefix="Perguntas norteadoras —")
    bullet(doc, "armadilhas comuns e o que o avaliador olha primeiro.", bold_prefix="Cuidados —")
    para(doc,
         "Símbolo de pergunta (?) em azul = pergunta para você responder pensando. Caixas amarelas = "
         "atenção/armadilha. Caixas verdes = material pronto. Caixas vermelhas = pendência (depende "
         "do notebook de Machine Learning, ainda não executado).")
    pagebreak(doc)

    # ===== PARTE I — ORDEM DE PRODUÇÃO =====
    h1(doc, "Parte I — A ordem certa de produção")
    para(doc,
         "O erro mais comum é escrever o artigo na ordem em que ele é lido (Introdução primeiro, "
         "Conclusão por último). O seu próprio plano de ação, baseado em Pereira & Galvão, prescreve "
         "o contrário. A lógica: você só sabe introduzir e concluir DEPOIS de saber o que encontrou.")
    map_table(doc, [
        ["1º", "Tabelas e Figuras", "Monte/organize todos os produtos visuais antes de escrever uma "
         "linha. Já temos a maioria — ver Parte II.", "FASE 1 do plano"],
        ["2º", "Materiais e Métodos", "Descreva o que foi FEITO. É a seção mais factual e a mais "
         "alimentada pelos nossos notebooks. Comece por aqui.", "FASE 2"],
        ["3º", "Resultados", "Reporte os achados SEM interpretar. Só os números e o que as "
         "tabelas/figuras mostram.", "FASE 3"],
        ["4º", "Discussão / Conclusão", "Agora sim interprete: o que os achados significam, "
         "comparados à literatura. Aqui entram nossas 5 discussões propostas.", "FASE 4"],
        ["5º", "Introdução", "Escrita quase no fim: agora você sabe exatamente para onde o artigo "
         "caminhou, então a introdução fica coerente com a conclusão.", "FASE 5"],
        ["6º", "Resumo / Abstract", "POR ÚLTIMO. É a miniatura do artigo pronto. Impossível escrever "
         "bem antes de tudo estar fechado.", "FASE 6"],
        ["7º", "Checklist final", "Rodar os 29 critérios do Prof. Zárate — ver Parte V.", "FASE 7"],
    ], headers=["Ordem", "Seção", "Por quê / o que fazer", "Plano"],
       col_widths=[1.3, 3.2, 8.5, 2.5])
    callout(doc, "A regra de ouro do Prof. Zárate",
            ["O avaliador verifica primeiro se o OBJETIVO (final da Introdução) e a CONCLUSÃO "
             "combinam exatamente. Escreva os dois lado a lado e confira se um é a resposta do outro.",
             "Corolário: como você escreve a Introdução no 5º lugar, deixe o objetivo provisório "
             "anotado desde já, mas só finalize-o quando a Conclusão estiver pronta."])
    pagebreak(doc)

    # ===== PARTE II — MAPA GERAL =====
    h1(doc, "Parte II — Mapa: o que já temos para cada seção")
    para(doc,
         "Visão de helicóptero. Esta tabela cruza cada seção do seu modelo de artigo com o material "
         "que produzimos nos notebooks 02 (análise exploratória e bivariada), 03 (pré-processamento "
         "— artrite pura) e 03B (pré-processamento — artrite com comorbidades), além do documento "
         "Resultados_Consolidados.")
    map_table(doc, [
        ["Resumo / Abstract", "Nada ainda — escrito por último", "PENDENTE"],
        ["1. Introdução", "Prevalência calculada por nós (17,71%); plano tem contexto e referências",
         "PARCIAL"],
        ["2. Trabalhos Relacionados", "Plano lista 8 referências (Tabela 26 do plano)", "PARCIAL"],
        ["3.1 Materiais (PNS 2019)", "Descrição das 4 bases; plano tem texto-padrão", "PRONTO"],
        ["3.2 Etapa 1 — Problema", "Estrutura PICOS completa (plano)", "PRONTO"],
        ["3.3 Etapa 2 — CAPTO", "Tabela 18 do plano (mapa CAPTO gigante por dimensão)", "PRONTO"],
        ["3.4 Etapa 3 — Seleção de atributos", "Listas tipadas dos notebooks 03/03B + Tabela 18 plano",
         "PRONTO"],
        ["3.5 Etapa 4 — Dados ausentes", "Notebooks 03/03B: skip patterns + auditoria de missing",
         "PRONTO"],
        ["3.6 Etapa 5 — Outliers", "Notebooks 03/03B: log de outliers (IQR 3x por classe)", "PRONTO"],
        ["3.7 Etapa 6 — Preparação", "Notebooks 03/03B: fusão (IMC, escores) + discretização", "PRONTO"],
        ["3.8 Etapa 7 — Seleção (entropia)", "Não feito — pertence ao notebook de ML", "PENDENTE"],
        ["3.9 Etapa 8 — Balanceamento", "Razões conhecidas (8,77:1 e 1,08:1); RUS é no notebook de ML",
         "PARCIAL"],
        ["3.10 Etapa 9 — Modelos", "Não feito — notebook de ML", "PENDENTE"],
        ["3.11 Etapa 10 — Treino/teste", "Não feito — notebook de ML", "PENDENTE"],
        ["4. Resultados", "EDA/bivariada PRONTA (Tab 2A-3B); ML PENDENTE (Tab 4-5, Gráf 1-2)",
         "PARCIAL"],
        ["5. Conclusão / Discussão", "5 discussões propostas + limitações (Resultados_Consolidados)",
         "PARCIAL"],
    ], headers=["Seção do artigo", "Material disponível", "Status"],
       col_widths=[4.5, 8.5, 2.5])
    callout(doc, "Leitura do mapa",
            ["Tudo que é PRÉ-PROCESSAMENTO e ANÁLISE DESCRITIVA já está pronto — é cerca de 60% do "
             "artigo. O que falta (PENDENTE) é toda a parte de Machine Learning: seleção por entropia, "
             "balanceamento RUS, treino dos algoritmos, métricas e feature importance. Essa parte "
             "depende do próximo notebook (04).",
             "Estratégia recomendada: escreva AGORA tudo que está PRONTO (Materiais e Métodos até a "
             "Etapa 7, e a caracterização da amostra). Deixe os slots de ML claramente marcados como "
             "[A PREENCHER APÓS NOTEBOOK 04]. Você terá ~60% do artigo redigido antes mesmo de "
             "treinar o primeiro modelo."], cor="D5E8F0")
    pagebreak(doc)

    # ===== PARTE III — GUIA SEÇÃO POR SEÇÃO =====
    h1(doc, "Parte III — Guia seção por seção")
    para(doc,
         "O coração do guia. Trabalhe na ORDEM DA PARTE I (Métodos primeiro), não na ordem abaixo. "
         "A ordem abaixo segue a numeração do artigo apenas para facilitar a consulta.")

    # --- 3. MATERIAIS E MÉTODOS (escrever 1º) ---
    h2(doc, "3. Materiais e Métodos  —  ESCREVER PRIMEIRO")
    para(doc,
         "Esta é a seção mais alimentada pelo nosso trabalho e a mais factual. Não há interpretação "
         "aqui — você apenas descreve, com precisão, o que foi feito. O Prof. Zárate exige que "
         "Materiais e Métodos sejam separados (item 7) e que as 10 etapas do pipeline apareçam "
         "individualizadas (itens 11-22).")

    h3(doc, "3.1 Materiais — a base PNS 2019 e as quatro bases derivadas")
    para(doc, "O que o gabarito exige:", bold=True)
    bullet(doc, "Item 9: base de dados descrita com detalhes.")
    bullet(doc, "Item 10: análise descritiva da base apresentada.")
    para(doc, "Material que já temos:", bold=True)
    bullet(doc, "Descrição da PNS 2019 (293.726 registros, 1.088 atributos) — texto-padrão no plano.")
    bullet(doc, "Tabela das 4 bases derivadas e seus N — ver Tabela 2.2 do Resultados_Consolidados.")
    bullet(doc, "Filtros de integridade V0015 e M001 — descritos nos notebooks 03/03B.")
    para(doc, "Perguntas norteadoras:", bold=True)
    question(doc, "Como você explica, em 3-4 frases, o que é a PNS 2019 e por que ela serve ao "
                  "seu objetivo? (Dica: cobertura nacional, dados de saúde autorreferidos, domínio público.)")
    question(doc, "Por que você construiu QUATRO bases e não apenas duas? O que cada uma representa "
                  "clinicamente? (Geral, Saudáveis, Artrite com comorbidades, Artrite pura.)")
    question(doc, "O que são os filtros V0015=1 e M001=1 e por que sem eles a amostra ficaria "
                  "contaminada com entrevistas incompletas?")
    question(doc, "Qual o N final de cada grupo nos dois desenhos de estudo? (494 vs 4.332 no "
                  "desenho puro; 4.025 vs 4.332 no desenho com comorbidades.)")

    h3(doc, "3.2 Etapa 1 — Entendimento do problema (escopo)")
    para(doc, "Material que já temos: a estrutura PICOS completa está no plano de ação (Tabela 13). "
              "Use-a literalmente.", bold=False)
    para(doc, "Perguntas norteadoras:", bold=True)
    question(doc, "Em uma frase: qual é o problema que justifica este estudo existir?")
    question(doc, "Por que a artrite em idosos brasileiros é um problema de saúde pública relevante "
                  "o suficiente para merecer um artigo?")

    h3(doc, "3.3 Etapa 2 — Entendimento do domínio e modelagem conceitual (CAPTO)")
    para(doc, "Material que já temos: a Tabela 18 do plano de ação é o mapa CAPTO inteiro, organizado "
              "por dimensão (hábitos de saúde, hábitos alimentares, etc.) com cada atributo, sua fonte "
              "científica e suas estatísticas. É a sua maior fonte para esta etapa.", bold=False)
    para(doc, "Perguntas norteadoras:", bold=True)
    question(doc, "Quais são as DIMENSÕES conceituais do seu problema? (Nutrição, atividade física, "
                  "sociodemografia, saúde física/mental, antropometria.)")
    question(doc, "Para cada dimensão, qual a justificativa científica de tê-la incluído? Que "
                  "referência sustenta cada uma? (O plano já lista: Malta 2021 para atividade física, "
                  "Louzada/Claro para dieta, etc.)")
    question(doc, "Você precisa inserir a FIGURA do modelo conceitual (Figura 1 do template). Você "
                  "já tem esse mapa desenhado? Se não, é uma pendência a resolver — pode ser um "
                  "diagrama simples ligando as dimensões ao atributo-alvo (artrite).")
    callout(doc, "Pendência identificada",
            "O template pede a Figura 1 (Modelo Conceitual). Não temos essa figura nos notebooks — "
            "ela é conceitual, não gerada por código. Você precisará desenhá-la (PowerPoint, draw.io, "
            "ou à mão) a partir da Tabela 18 do plano.", cor="FFF3CD")

    h3(doc, "3.4 Etapa 3 — Seleção conceitual de atributos")
    para(doc, "O que o gabarito exige: Item 14 — seleção conceitual de atributos COM TABELA "
              "(é a Tabela 1 do artigo).", bold=False)
    para(doc, "Material que já temos:", bold=True)
    bullet(doc, "As listas tipadas dos notebooks 03/03B: VARS_CONTINUAS, VARS_DISCRETAS, "
                "VARS_NOMINAIS, VARS_ORDINAIS — cada variável com seu código PNS e descrição.")
    bullet(doc, "A Tabela 18 do plano dá a 'razão de seleção' de cada atributo.")
    para(doc, "Perguntas norteadoras:", bold=True)
    question(doc, "Quantos atributos você selecionou no total? O template pede no mínimo 10 + o alvo. "
                  "(Você tem bem mais — precisa decidir se reporta todos ou os principais.)")
    question(doc, "Para a coluna 'Razão de Seleção' da Tabela 1: por que CADA atributo foi escolhido? "
                  "O que ele captura do domínio?")
    question(doc, "Qual é o atributo-alvo e seu código PNS? (Q079 — diagnóstico de artrite/reumatismo.)")

    h3(doc, "3.5 Etapa 4 — Tratamento de dados ausentes e vazios")
    callout(doc, "Este é um dos seus pontos fortes — explore-o bem",
            ["A maioria dos trabalhos trata missing de forma genérica. Você fez algo mais "
             "sofisticado: distinguiu missing ESTRUTURAL (skip patterns do questionário) de missing "
             "ALEATÓRIO real. Isso é um diferencial metodológico — reporte com orgulho."], cor="D5EED5")
    para(doc, "Material que já temos:", bold=True)
    bullet(doc, "Tabela de skip patterns resolvidos (etapa2_5_skip_patterns.csv) — nos dois desenhos.")
    bullet(doc, "Auditoria de missing com critério de exclusão >75% (etapa3_auditoria_missing.csv).")
    bullet(doc, "Figura do heatmap de missing por variável e grupo — ver Figura 4.1 do "
                "Resultados_Consolidados.")
    bullet(doc, "Método de imputação: média/moda por classe (etapa5_log_imputacao.csv).")
    para(doc, "Perguntas norteadoras:", bold=True)
    question(doc, "O que é um 'skip pattern' do questionário da PNS? Dê o exemplo mais didático "
                  "(quem responde 'não pratico exercício' não é perguntado 'quantos dias por semana').")
    question(doc, "Por que tratar esse vazio como missing aleatório e imputá-lo pela média seria um "
                  "ERRO metodológico?")
    question(doc, "Quantas variáveis foram excluídas por terem mais de 75% de dados ausentes reais? "
                  "Quais eram, e por que isso é uma limitação? (Renda, autoavaliação de saúde, "
                  "escolaridade — perda da dimensão socioeconômica.)")
    question(doc, "Por que a imputação foi feita pela média/moda DA PRÓPRIA CLASSE e não pela média "
                  "global? E qual o risco disso para a fase de ML? (Vazamento — refazer dentro dos "
                  "folds de validação cruzada.)")

    h3(doc, "3.6 Etapa 5 — Análise e remoção de outliers")
    para(doc, "Material que já temos: log de outliers dos dois desenhos (etapa4_log_outliers.csv) e "
              "a Figura 4.2 do Resultados_Consolidados (boxplots antes/depois).", bold=False)
    para(doc, "Perguntas norteadoras:", bold=True)
    question(doc, "Qual método de detecção de outliers você usou e por quê? (IQR — intervalo "
                  "interquartil.)")
    question(doc, "Você usou 3×IQR em vez do 1,5×IQR padrão. Por quê? (Mais conservador — só remove "
                  "valores extremos que são prováveis erros de medição, não variação biológica.)")
    question(doc, "Por que o IQR foi calculado SEPARADAMENTE por classe (artrite vs. saudável)? "
                  "O que aconteceria se fosse calculado misturando os dois grupos?")
    question(doc, "Quantos outliers foram substituídos no total, e em quais variáveis se "
                  "concentraram?")
    callout(doc, "Atenção — divergência com o template",
            "O texto-padrão do template (Tabela 24 do modelo) menciona '1,5×IQR'. Você usou 3×IQR. "
            "NÃO copie o texto-padrão cegamente — ajuste para 3×IQR e justifique a escolha. "
            "Inconsistência entre o que você fez e o que escreveu é exatamente o que o avaliador "
            "procura.", cor="FFF3CD")

    h3(doc, "3.7 Etapa 6 — Preparação dos dados (fusão, transformação, discretização)")
    para(doc, "O que o gabarito exige: Item 17 — combinação, fusão e discretização. Item 18 — "
              "apresentar o conjunto de dados final. É a Tabela 2 do artigo.", bold=False)
    para(doc, "Material que já temos:", bold=True)
    bullet(doc, "Fusão: IMC (de peso+altura), Escore Inflamatório e Escore Saudável (de variáveis "
                "alimentares), Razão Inflamatório/Saudável.")
    bullet(doc, "Discretização: 4 variáveis — FxEtaria_cat, IMC_cat, AtivFisica_cat, EscInfla_cat — "
                "com a metodologia já documentada (pd.cut por limiares OMS e pd.qcut por quartis).")
    bullet(doc, "Encoding: Label Encoding (ordinais) + One-Hot Encoding (nominais).")
    bullet(doc, "Datasets finais: 4.826 × 49 (puro) e 8.357 × 57 (comorbidades).")
    para(doc, "Perguntas norteadoras:", bold=True)
    question(doc, "Como o IMC foi criado e por que combinar peso e altura num único atributo é "
                  "melhor do que mantê-los separados?")
    question(doc, "O que são os escores Inflamatório e Saudável? Que hipótese clínica eles "
                  "operacionalizam?")
    question(doc, "Para a Tabela 2 do artigo: para cada variável discretizada, qual a regra de "
                  "transformação e quais as categorias resultantes? (Você já tem isso documentado — "
                  "ver o guia de discretização que conversamos.)")
    question(doc, "Por que a idade e o IMC foram discretizados por limiares FIXOS (OMS), enquanto o "
                  "escore inflamatório foi discretizado por QUARTIS? O que justifica métodos "
                  "diferentes?")
    question(doc, "Qual o N e o número de atributos do dataset final? Você vai reportar UM desenho "
                  "ou os DOIS (puro e comorbidades)?")

    h3(doc, "3.8 a 3.11 — Etapas 7 a 10 (Seleção por entropia, Balanceamento, Modelos, Treino)")
    callout(doc, "PENDENTE — depende do notebook de Machine Learning (04)",
            ["Estas quatro etapas ainda não foram executadas. Elas pertencem ao próximo notebook. "
             "O que você JÁ pode adiantar agora:",
             "• Etapa 8 (Balanceamento): você já SABE as razões de desbalanceamento — 8,77:1 no "
             "desenho puro e 1,08:1 no desenho com comorbidades. Pode descrever que o desenho puro "
             "exigirá Random Under Sampler e o desenho com comorbidades praticamente não.",
             "• Etapas 7, 9 e 10: deixe os slots marcados como [A PREENCHER APÓS NOTEBOOK 04] e siga "
             "para as outras seções. O plano de ação (Tabela 19) já tem o texto-base de quais "
             "algoritmos e validação você usará."], cor="F5D5D5")
    para(doc, "Perguntas norteadoras (para já ir pensando):", bold=True)
    question(doc, "Por que o desbalanceamento do desenho 'artrite pura' (8,77:1) é um problema, e "
                  "como o Random Under Sampler o resolve? Qual o custo de usá-lo?")
    question(doc, "Quais 4 algoritmos você usará e por que essa mistura de 'caixa-branca' (Árvore, "
                  "Naive Bayes) e 'caixa-preta' (Random Forest, AdaBoost)? (Ver Loyola-González 2019, "
                  "já no template.)")

    pagebreak(doc)

    # --- 4. RESULTADOS ---
    h2(doc, "4. Experimentos e Análise dos Resultados  —  ESCREVER EM TERCEIRO")
    callout(doc, "Decisão estrutural importante que você precisa tomar",
            ["O template do Prof. Zárate foi desenhado esperando que a Seção 4 contenha SOMENTE "
             "resultados de Machine Learning (Tabelas 4-5, Gráficos 1-2). Mas você produziu uma "
             "análise exploratória e bivariada rica que não tem 'slot' óbvio no template.",
             "Recomendação: a sua análise descritiva/bivariada (prevalência, comparação de grupos, "
             "testes de hipótese) entra como uma SUBSEÇÃO inicial — pode ser '4.1 Caracterização da "
             "Amostra' — antes dos resultados de ML. É o equivalente à 'Tabela 1' de artigos "
             "epidemiológicos. Discuta essa organização com o orientador, mas tê-la fortalece muito "
             "o artigo: mostra que você entendeu os dados antes de modelar."], cor="D5E8F0")

    h3(doc, "4.1 (sugerido) Caracterização da amostra — análise descritiva e bivariada")
    para(doc, "Material que já temos — TUDO PRONTO:", bold=True)
    bullet(doc, "Prevalência de artrite por subgrupo — Tabela 3.1 e Figura 3.1 do Resultados_Consolidados.")
    bullet(doc, "Comparação das 4 bases — Tabela 3.2 e Figura 3.2.")
    bullet(doc, "Características quantitativas (artrite pura vs. saudáveis) — Tabela 3.3 e Figura 3.3.")
    bullet(doc, "Características qualitativas — Tabela 3.4 e Figura 3.4.")
    bullet(doc, "Análise bivariada completa com testes de hipótese — Tabela 3.5.")
    bullet(doc, "Caracterização da multimorbidade — Tabela 4.3 e Figura 4.3.")
    para(doc, "Perguntas norteadoras (lembre-se: Resultados NÃO interpreta, só reporta):", bold=True)
    question(doc, "Qual a prevalência global de artrite na sua amostra? Como ela varia por sexo, "
                  "idade e escolaridade? (Reporte os números — a interpretação vai para a Discussão.)")
    question(doc, "Quais variáveis diferiram significativamente entre idosos com artrite e saudáveis? "
                  "Liste-as com seus p-valores. (Foram 6: sexo, autoavaliação de saúde, plano de "
                  "saúde, IMC, frutas, refrigerante.)")
    question(doc, "Quantas comorbidades, em média, tem um idoso com artrite? Quais as 3 mais "
                  "prevalentes? (Hipertensão 65,3%, colesterol 39,8%, diabetes 21,9%.)")
    question(doc, "Como você apresenta isso sem repetir no texto o que já está na tabela? (Regra de "
                  "ouro: o texto destaca o achado principal; a tabela tem o detalhe.)")

    h3(doc, "4.2 a 4.4 — Resultados de Machine Learning")
    callout(doc, "PENDENTE — notebook de Machine Learning (04)",
            ["Tabela 4 (treinamento com IC 95%), Tabela 5 (teste), Gráfico 1 (comparação de "
             "algoritmos), Gráfico 2 (feature importance) e as regras da árvore de decisão — todos "
             "dependem do notebook 04. Deixe esses slots marcados e volte a eles depois.",
             "Mas você já pode escrever a ESTRUTURA da seção e as frases de ligação, deixando só os "
             "números em branco."], cor="F5D5D5")
    para(doc, "Perguntas norteadoras (para depois do notebook 04):", bold=True)
    question(doc, "Qual algoritmo teve o melhor F1-Score? A escolha do 'melhor modelo' considera só "
                  "a métrica ou também a interpretabilidade?")
    question(doc, "Quais foram os atributos de maior importância? Eles confirmam ou contrariam a "
                  "hipótese nutricional inicial?")
    question(doc, "Que regras a Árvore de Decisão extraiu? Alguma é clinicamente interpretável?")

    pagebreak(doc)

    # --- 5. CONCLUSÃO / DISCUSSÃO ---
    h2(doc, "5. Conclusão e Trabalhos Futuros  —  ESCREVER EM QUARTO")
    para(doc, "O que o gabarito exige: Item 27 (observações dos resultados e da metodologia), "
              "Item 28 (limitações), Item 29 (trabalhos futuros). Regra de ouro: a conclusão DEVE "
              "combinar com o objetivo da introdução.", bold=False)
    para(doc, "Material que já temos — as 5 discussões propostas no Resultados_Consolidados:", bold=True)
    bullet(doc, "O paradoxo nutricional e a causalidade reversa (artríticos comem mais frutas — "
                "contraintuitivo).")
    bullet(doc, "O IMC como elo entre artrite e comorbidades.")
    bullet(doc, "O perfil sociodemográfico: a artrite tem gênero e idade.")
    bullet(doc, "O dilema metodológico dos dois desenhos de estudo.")
    bullet(doc, "A multimorbidade como achado clínico central.")
    para(doc, "E a lista de 6 limitações já redigida no mesmo documento.")
    para(doc, "Perguntas norteadoras:", bold=True)
    question(doc, "Em 2-3 frases: qual é o achado principal do seu trabalho? (Cuidado: você só "
                  "saberá 100% depois do ML, mas a parte descritiva já permite uma resposta parcial.)")
    question(doc, "O seu objetivo declarado na Introdução e a sua conclusão dizem a MESMA coisa? "
                  "Escreva os dois lado a lado e confira.")
    question(doc, "Quais são as limitações honestas do trabalho? (Já temos 6 listadas — quais se "
                  "aplicam de fato? Design transversal, dados autorreferidos, vazamento de "
                  "delineamento no desenho B, etc.)")
    question(doc, "Que trabalhos futuros fazem sentido? (Dados longitudinais, biomarcadores "
                  "laboratoriais, replicar com PNS 2024.)")
    callout(doc, "Cuidado — a Discussão não pode passar de 1/3 do artigo",
            "É um critério explícito do Prof. Zárate. Se a sua Discussão estiver muito longa, "
            "provavelmente você está repetindo resultados ou interpretando demais. Seja cirúrgico.",
            cor="FFF3CD")

    pagebreak(doc)

    # --- 1. INTRODUÇÃO ---
    h2(doc, "1. Introdução  —  ESCREVER EM QUINTO")
    para(doc, "O que o gabarito exige: Item 1 (contexto epidemiológico), Item 2 (trabalhos "
              "relacionados citados brevemente), Item 3 (problema + fonte + técnica), Item 4 "
              "(estrutura do artigo no último parágrafo). Regra de ouro: o objetivo fica no FINAL.",
         bold=False)
    para(doc, "Material que já temos:", bold=True)
    bullet(doc, "Um dado epidemiológico CALCULADO POR VOCÊ: prevalência de 17,71% de artrite entre "
                "idosos. Isso é mais forte que citar literatura — é seu próprio resultado.")
    bullet(doc, "Contexto e estado da arte — Tabelas 2, 3 e 24 do plano de ação.")
    bullet(doc, "Objetivo já redigido — Tabela 7 do plano.")
    para(doc, "Perguntas norteadoras:", bold=True)
    question(doc, "Qual dado de impacto abre o artigo? (Sugestão: combine a prevalência que VOCÊ "
                  "calculou com dados da OMS/literatura.)")
    question(doc, "Em 2-3 frases, que trabalhos anteriores você menciona — e por que o seu é "
                  "diferente? (Não é a seção de Trabalhos Relacionados ainda, é só uma menção.)")
    question(doc, "O último parágrafo descreve a estrutura do artigo por seções — o template já tem "
                  "esse parágrafo pronto. Ele ainda bate com a sua estrutura real?")
    question(doc, "O seu objetivo final combina exatamente com a conclusão que você escreveu na "
                  "etapa anterior?")

    # --- 2. TRABALHOS RELACIONADOS ---
    h2(doc, "2. Trabalhos Relacionados  —  pode escrever junto com os Métodos")
    para(doc, "O que o gabarito exige: Item 5 (mínimo 6 trabalhos), Item 6 (apontar a diferença do "
              "seu trabalho).", bold=False)
    para(doc, "Material que já temos: o plano de ação (Tabela 26) lista 8 referências com o motivo "
              "de usar cada uma — Cancella & Zárate 2025, Melo et al. 2024, Silva & Zárate 2024, "
              "Gonçalves & Zárate 2024, Zárate et al. 2023 (CAPTO), Xue et al. 2020, IBGE 2020, "
              "Malta et al. 2021.", bold=False)
    para(doc, "Perguntas norteadoras:", bold=True)
    question(doc, "Você tem pelo menos 6 trabalhos? Quais aplicam ML sobre a PNS? Quais tratam "
                  "especificamente de artrite?")
    question(doc, "Para cada trabalho: contexto → técnica → base de dados → resultado → conclusão. "
                  "Consegue resumir cada um nesse padrão de 1-2 frases?")
    question(doc, "Qual a diferença do SEU trabalho? (Foco específico em artrite na população idosa "
                  "brasileira; dois desenhos de estudo paralelos; método CAPTO.)")

    # --- RESUMO ---
    h2(doc, "Resumo / Abstract  —  ESCREVER POR ÚLTIMO")
    para(doc, "Estrutura obrigatória: Objetivo + Método + Resultados + Conclusão, em 150-250 "
              "palavras. O JHI exige versões em português, inglês e espanhol.", bold=False)
    para(doc, "Perguntas norteadoras:", bold=True)
    question(doc, "Consegue resumir o objetivo em 1-2 frases? E o método? E o resultado principal? "
                  "E a conclusão?")
    question(doc, "O objetivo e a conclusão do resumo combinam entre si? (Mesma verificação da regra "
                  "de ouro.)")
    callout(doc, "Por que o resumo é o último",
            "O resumo é uma miniatura fiel do artigo PRONTO. Escrevê-lo antes te obriga a adivinhar "
            "resultados que você ainda não tem. Deixe-o por último — quando ele virar só um trabalho "
            "de síntese, não de adivinhação.", cor="D5E8F0")

    pagebreak(doc)

    # ===== PARTE IV — STATUS =====
    h1(doc, "Parte IV — O que está pronto e o que falta")
    para(doc, "Inventário honesto do estado atual do projeto, para você priorizar.")
    h2(doc, "Já produzido (notebooks 02, 03 e 03B)")
    status_table(doc, [
        ["PRONTO", "Análise exploratória e bivariada completa",
         "Notebook 02 — tabelas 2A-3B e figuras 2A-3B"],
        ["PRONTO", "Pré-processamento — desenho artrite pura",
         "Notebook 03 — skip patterns, missing, outliers, fusão, discretização, encoding"],
        ["PRONTO", "Pré-processamento — desenho artrite com comorbidades",
         "Notebook 03B — idem + análise de multimorbidade"],
        ["PRONTO", "Caracterização das comorbidades",
         "Notebook 03B — tabela e figura de prevalência por doença"],
        ["PRONTO", "Discussões metodológicas e limitações",
         "Resultados_Consolidados — Seções 5 e 6"],
        ["PRONTO", "Metodologia de discretização documentada",
         "Conversa do chat — pd.cut por OMS e pd.qcut por quartis"],
    ])
    h2(doc, "Ainda pendente (notebook de Machine Learning — 04)")
    status_table(doc, [
        ["PENDENTE", "Seleção de atributos por entropia de Shannon",
         "Etapa 7 do pipeline — fazer no notebook 04"],
        ["PENDENTE", "Balanceamento com Random Under Sampler",
         "Etapa 8 — aplicar dentro dos folds de validação cruzada"],
        ["PENDENTE", "Treino dos 4 algoritmos + RandomizedSearchCV",
         "Etapa 9-10 — Árvore, Naive Bayes, Random Forest, AdaBoost"],
        ["PENDENTE", "Métricas com IC 95% (Tabelas 4 e 5 do artigo)",
         "Validação cruzada 10-fold + holdout"],
        ["PENDENTE", "Feature Importance (Gráfico 2 do artigo)",
         "scikit-learn — modelos baseados em árvore"],
        ["PENDENTE", "Regras da Árvore de Decisão",
         "Extração de regras interpretáveis"],
        ["PENDENTE", "Figura 1 — Modelo Conceitual CAPTO",
         "NÃO é código — desenhar manualmente a partir da Tabela 18 do plano"],
    ])
    callout(doc, "Conclusão do inventário",
            ["Você pode redigir AGORA, sem depender de mais nada: toda a Seção 3 (Materiais e "
             "Métodos) até a Etapa 6, e a subseção 4.1 (Caracterização da Amostra). Isso é "
             "aproximadamente 55-60% do texto do artigo.",
             "O caminho crítico para fechar o artigo é o notebook 04 (Machine Learning). Ele "
             "destrava as Etapas 7-10 da metodologia e toda a parte 4.2-4.4 dos resultados."],
            cor="D5E8F0")

    pagebreak(doc)

    # ===== PARTE V — CHECKLIST =====
    h1(doc, "Parte V — Checklist do Prof. Zárate (29 critérios) com status atual")
    para(doc, "Os 29 critérios de avaliação do template, anotados com o estado atual: o que o "
              "material já produzido cobre, o que está parcial e o que ainda depende do notebook 04.")
    map_table(doc, [
        ["1", "Domínio contextualizado com dados de impacto", "PARCIAL — temos a prevalência calculada"],
        ["2", "Trabalhos relacionados citados na Introdução", "PARCIAL — plano lista 8 referências"],
        ["3", "Problema, fonte e técnica mencionados", "PRONTO — definido no plano"],
        ["4", "Estrutura por seções no fim da Introdução", "PRONTO — template já tem o parágrafo"],
        ["5", "Mínimo 6 trabalhos relacionados", "PARCIAL — plano tem 8, falta redigir"],
        ["6", "Diferença do trabalho proposto apontada", "PARCIAL — dois desenhos é o diferencial"],
        ["7", "Materiais e Métodos separados", "A FAZER na redação"],
        ["9", "Base de dados descrita com detalhes", "PRONTO — descrição das 4 bases"],
        ["10", "Análise descritiva da base apresentada", "PRONTO — notebook 02 inteiro"],
        ["11", "Etapas do pipeline separadas", "PRONTO até etapa 6; resto pendente"],
        ["12", "Etapa 1 — escopo do problema", "PRONTO — PICOS no plano"],
        ["13", "Etapa 2 — domínio e CAPTO", "PRONTO — Tabela 18 do plano (falta Figura 1)"],
        ["14", "Etapa 3 — seleção de atributos com tabela", "PRONTO — listas dos notebooks 03/03B"],
        ["15", "Etapa 4 — dados ausentes", "PRONTO — skip patterns + auditoria de missing"],
        ["16", "Etapa 5 — outliers", "PRONTO — log de outliers IQR 3x"],
        ["17", "Etapa 6 — preparação (fusão, discretização)", "PRONTO — IMC, escores, 4 var. discretizadas"],
        ["18", "Conjunto de dados final apresentado", "PRONTO — 4.826x49 e 8.357x57"],
        ["19", "Etapa 7 — seleção por entropia", "PENDENTE — notebook 04"],
        ["20", "Etapa 8 — balanceamento", "PARCIAL — razões conhecidas; RUS no notebook 04"],
        ["21", "Etapa 9 — modelos e parametrização", "PENDENTE — notebook 04"],
        ["22", "Etapa 10 — treino e teste", "PENDENTE — notebook 04"],
        ["23", "Seção de experimentos e resultados", "PARCIAL — EDA pronta; ML pendente"],
        ["24", "Medidas de desempenho e IC 95%", "PENDENTE — notebook 04"],
        ["25", "Modelo final apresentado", "PENDENTE — notebook 04"],
        ["26", "Análise do conhecimento extraído (regras)", "PENDENTE — notebook 04"],
        ["27", "Conclusões com observações", "PARCIAL — 5 discussões propostas prontas"],
        ["28", "Limitações apontadas", "PRONTO — 6 limitações redigidas"],
        ["29", "Trabalhos futuros apontados", "PRONTO — listados no Resultados_Consolidados"],
    ], headers=["Nº", "Critério", "Status atual"], col_widths=[1.0, 8.0, 6.5])
    callout(doc, "Como ler este checklist",
            ["Conte: dos 29 critérios, a maioria já está PRONTA ou PARCIAL. Os 6-7 PENDENTES "
             "concentram-se todos no notebook de Machine Learning (critérios 19, 21, 22, 24, 25, 26) "
             "mais a Figura 1 conceitual.",
             "Use esta tabela como painel de controle: a cada sessão de escrita, atualize um status "
             "de PARCIAL para PRONTO. Quando todos estiverem PRONTO, o artigo está pronto para a "
             "revisão do orientador."], cor="D5E8F0")

    # Rodapé
    doc.add_paragraph()
    fim = doc.add_paragraph(); fim.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = fim.add_run("— Fim do guia de redação —\n"
                     "Trabalhe na ordem da Parte I. Responda às perguntas norteadoras com suas "
                     "palavras. O artigo nascerá dessas respostas.")
    rf.italic = True; rf.font.size = Pt(9); rf.font.color.rgb = CINZA

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(f"[OK] Guia gerado: {OUT_PATH}")
    print(f"     {len(doc.paragraphs)} paragrafos, {len(doc.tables)} tabelas")


if __name__ == "__main__":
    build()
