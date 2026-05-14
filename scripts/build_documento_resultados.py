# -*- coding: utf-8 -*-
"""Gera o documento Word consolidado de resultados para apoiar o artigo científico.

Reúne resultados dos notebooks 02 (análise exploratória/bivariada), 03 (pré-
processamento — artrite pura) e 03B (pré-processamento — artrite com comorbidades),
inserindo as tabelas e figuras nas seções de análise, com discussões propostas.

Uso (rodar da raiz do projeto):
    python scripts/build_documento_resultados.py

Gera: Documentos_organizacao/Resultados_Consolidados_PNS2019_Artrite.docx
"""
from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
RES_EDA   = ROOT / "data" / "results" / "eda"
RES_PURO  = ROOT / "data" / "results" / "preprocessing"
RES_COMOR = ROOT / "data" / "results" / "preprocessing_comorbidades"
OUT_PATH  = ROOT / "Documentos_organizacao" / "Resultados_Consolidados_PNS2019_Artrite.docx"

# Paleta
AZUL_ESCURO = RGBColor(0x1F, 0x37, 0x64)
AZUL_MEDIO  = RGBColor(0x2E, 0x75, 0xB6)
CINZA       = RGBColor(0x59, 0x59, 0x59)
VERMELHO    = RGBColor(0xC0, 0x39, 0x2B)
VERDE       = RGBColor(0x27, 0xAE, 0x60)


# ─────────────────────────────────────────────────────────────────────────────
# Helpers de formatação
# ─────────────────────────────────────────────────────────────────────────────
def setup_styles(doc: Document) -> None:
    """Define fonte padrão e estilos de cabeçalho."""
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    for lvl, size, color in [(1, 18, AZUL_ESCURO), (2, 14, AZUL_MEDIO), (3, 12, AZUL_MEDIO)]:
        st = doc.styles[f"Heading {lvl}"]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color


def add_cover(doc: Document) -> None:
    """Capa do documento."""
    for _ in range(3):
        doc.add_paragraph()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Caracterização de Idosos com Artrite e Reumatismo\nno Brasil — PNS 2019")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = AZUL_ESCURO

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Documento consolidado de resultados — Análise Exploratória, "
                     "Bivariada e Pré-Processamento")
    rs.italic = True
    rs.font.size = Pt(13)
    rs.font.color.rgb = CINZA

    doc.add_paragraph()
    linha = doc.add_paragraph()
    linha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    linha.add_run("_" * 55).font.color.rgb = AZUL_MEDIO

    for _ in range(2):
        doc.add_paragraph()

    info = [
        ("Pesquisador", "Pedro Dias Soares"),
        ("Orientador", "Prof. Dr. Luis Enrique Zárate — PUC Minas"),
        ("Linha de pesquisa", "Saúde Pública · Mineração de Dados · Aprendizado de Máquina"),
        ("Periódico-alvo", "Journal of Health Informatics (JHI / SBIS)"),
        ("Base de dados", "Pesquisa Nacional de Saúde 2019 (IBGE) — 293.726 registros"),
        ("Metodologia", "KDD · CAPTO · PICTOREA · STROBE"),
    ]
    for rotulo, valor in info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = p.add_run(f"{rotulo}:  ")
        rr.bold = True
        rr.font.size = Pt(11)
        rv = p.add_run(valor)
        rv.font.size = Pt(11)

    for _ in range(6):
        doc.add_paragraph()

    nota = doc.add_paragraph()
    nota.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn = nota.add_run("Gerado automaticamente a partir dos notebooks 02, 03 e 03B do repositório.\n"
                      "Os resultados aqui consolidados destinam-se a apoiar a redação das seções "
                      "de Resultados e Discussão do artigo científico.")
    rn.italic = True
    rn.font.size = Pt(9)
    rn.font.color.rgb = CINZA

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def h1(doc, txt):  return doc.add_heading(txt, level=1)
def h2(doc, txt):  return doc.add_heading(txt, level=2)
def h3(doc, txt):  return doc.add_heading(txt, level=3)


def para(doc, txt, italic=False, size=11, bold=False, justify=True):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(txt)
    r.italic = italic
    r.bold = bold
    r.font.size = Pt(size)
    return p


def bullet(doc, txt, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        rb = p.add_run(f"{bold_prefix} ")
        rb.bold = True
    p.add_run(txt)
    return p


def caption(doc, txt):
    p = doc.add_paragraph()
    r = p.add_run(txt)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = CINZA
    p.paragraph_format.space_after = Pt(10)
    return p


def add_figure(doc, path: Path, legenda: str, largura_cm: float = 15.5):
    """Insere figura centralizada com legenda."""
    if not path.exists():
        para(doc, f"[FIGURA NÃO ENCONTRADA: {path.name}]", italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(largura_cm))
    caption(doc, legenda)


def add_df_table(doc, df: pd.DataFrame, col_widths=None, header_color="2E75B6",
                 font_size=8, max_rows=None):
    """Converte um DataFrame em tabela formatada do Word."""
    if max_rows:
        df = df.head(max_rows)
    n_cols = len(df.columns)
    table = doc.add_table(rows=1, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Cabeçalho
    hdr = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr[i].text = str(col)
        for par in hdr[i].paragraphs:
            for run in par.runs:
                run.bold = True
                run.font.size = Pt(font_size)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade_cell(hdr[i], header_color)

    # Linhas
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = "" if pd.isna(val) else str(val)
            for par in cells[i].paragraphs:
                for run in par.runs:
                    run.font.size = Pt(font_size)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


def _shade_cell(cell, hex_color: str):
    """Aplica cor de fundo a uma célula."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): hex_color
    })
    tc_pr.append(shd)


def callout(doc, titulo: str, texto: str, cor="FFF3CD"):
    """Caixa de destaque (para discussões / avisos)."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    _shade_cell(cell, cor)
    cell.text = ""
    p1 = cell.paragraphs[0]
    r1 = p1.add_run(titulo)
    r1.bold = True
    r1.font.size = Pt(10.5)
    p2 = cell.add_paragraph()
    r2 = p2.add_run(texto)
    r2.font.size = Pt(10)
    doc.add_paragraph()


def load_csv(path: Path) -> pd.DataFrame:
    df = pd.read_csv(path)
    if df.columns[0].startswith("Unnamed"):
        df = df.drop(columns=df.columns[0])
    return df


# ─────────────────────────────────────────────────────────────────────────────
# Construção do documento
# ─────────────────────────────────────────────────────────────────────────────
def build():
    doc = Document()
    setup_styles(doc)

    # Margens
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    add_cover(doc)

    # Carrega relatórios JSON de pré-processamento
    rel_puro  = json.loads((RES_PURO  / "relatorio_preprocessamento.json").read_text(encoding="utf-8"))
    rel_comor = json.loads((RES_COMOR / "relatorio_preprocessamento.json").read_text(encoding="utf-8"))

    # ===========================================================================
    # SUMÁRIO
    # ===========================================================================
    h1(doc, "Sumário")
    sumario = [
        "1. Apresentação e objetivos",
        "2. Materiais e Métodos",
        "3. Resultados — Análise Exploratória e Bivariada",
        "4. Resultados — Pré-Processamento e os Dois Desenhos de Estudo",
        "5. Discussões Propostas para o Artigo",
        "6. Limitações Metodológicas",
        "7. Próximos Passos — Modelagem de Machine Learning",
        "Apêndice A — Índice de Tabelas e Figuras",
    ]
    for item in sumario:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ===========================================================================
    # 1. APRESENTAÇÃO E OBJETIVOS
    # ===========================================================================
    h1(doc, "1. Apresentação e Objetivos")

    para(doc,
         "Este documento consolida os resultados produzidos até o momento no projeto de "
         "mineração de dados sobre artrite e reumatismo em idosos brasileiros, com base nos "
         "microdados da Pesquisa Nacional de Saúde (PNS) 2019, conduzida pelo IBGE. O objetivo "
         "é reunir, em um único material organizado, todas as tabelas, figuras e análises "
         "estatísticas geradas pelos notebooks de análise exploratória e de pré-processamento, "
         "servindo como base para a redação das seções de Resultados e Discussão do artigo "
         "científico.")

    h2(doc, "1.1 Objetivo principal do estudo")
    para(doc,
         "Descrever as características nutricionais e sociodemográficas de idosos com artrite e "
         "reumatismo, com o propósito de descoberta de conhecimento em bases de dados por meio "
         "de modelos de aprendizado de máquina supervisionado aplicados à PNS 2019.")

    h2(doc, "1.2 Objetivos secundários")
    for txt in [
        "Mapear hábitos alimentares preditores (consumo de ultraprocessados, açúcar, alimentos "
        "in natura) extraídos da PNS 2019 via método CAPTO.",
        "Classificar os indivíduos entre perfil Saudável e perfil Portador de Artrite usando "
        "algoritmos de Aprendizado de Máquina Supervisionado (Árvore de Decisão, Floresta "
        "Aleatória, Naive Bayes, AdaBoost).",
        "Analisar a influência de covariáveis como atividade física, condições físicas/mentais "
        "e dados sociodemográficos e antropométricos na classificação do perfil de risco.",
        "Identificar os atributos de maior importância relativa para a classificação por meio "
        "de Feature Importance.",
    ]:
        bullet(doc, txt)

    h2(doc, "1.3 O que este documento cobre")
    para(doc,
         "O material aqui apresentado corresponde às fases já concluídas do processo KDD: a "
         "compreensão dos dados (análise exploratória e bivariada) e a preparação dos dados "
         "(pré-processamento completo). A fase de modelagem — treinamento dos algoritmos de "
         "classificação — será conduzida na etapa seguinte e está descrita na Seção 7.")

    callout(doc, "Destaque metodológico — dois desenhos de estudo paralelos",
            "Foram construídos dois datasets independentes para a modelagem: (A) Artrite Pura, "
            "que compara idosos com artrite isolada contra idosos saudáveis, e (B) Artrite com "
            "Comorbidades, que compara idosos com artrite (em qualquer combinação de doenças) "
            "contra idosos saudáveis. Ter os dois desenhos permite uma discussão mais rica sobre "
            "o trade-off entre pureza do contraste e balanceamento das classes — detalhado na "
            "Seção 4.", cor="D5E8F0")

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ===========================================================================
    # 2. MATERIAIS E MÉTODOS
    # ===========================================================================
    h1(doc, "2. Materiais e Métodos")

    h2(doc, "2.1 Delineamento do estudo (PICOS)")
    para(doc,
         "Trata-se de um estudo transversal retrospectivo baseado em microdados secundários da "
         "PNS 2019. A pergunta de pesquisa foi estruturada segundo o acrônimo PICOS:")
    picos = pd.DataFrame({
        "Componente": ["P — População", "I — Exposição", "C — Comparação",
                       "O — Desfecho", "S — Tipo de estudo"],
        "Definição no estudo": [
            "Idosos com 60 anos ou mais registrados na PNS 2019.",
            "Padrões nutricionais e variáveis sociodemográficas, físicas, mentais e antropométricas.",
            "Idosos saudáveis, sem diagnóstico de artrite nem de outras doenças crônicas relevantes.",
            "Classificação do perfil individual: Saudável vs. Portador de Artrite/Reumatismo.",
            "Estudo transversal com aplicação de processo KDD + aprendizado de máquina supervisionado.",
        ],
    })
    add_df_table(doc, picos, col_widths=[3.5, 12.0], font_size=9)
    caption(doc, "Tabela 2.1 — Estrutura PICOS da pergunta de pesquisa.")

    h2(doc, "2.2 Fonte de dados e construção das bases")
    para(doc,
         "A base original da PNS 2019 contém 293.726 registros e 1.088 atributos organizados "
         "em módulos temáticos. A partir dela foram construídas quatro bases de dados "
         "específicas por meio de filtros aplicados ao módulo Q (doenças crônicas):")
    bases = pd.DataFrame({
        "Base": ["População idosa geral", "Idosos saudáveis",
                 "Artrite com comorbidades", "Artrite pura"],
        "Critério": [
            "Idade ≥ 60 anos.",
            "Idade ≥ 60 anos e resposta 'Não' para todas as 14 doenças crônicas do módulo Q.",
            "Idade ≥ 60 anos e Q079 (artrite) = 'Sim', em qualquer combinação de comorbidades.",
            "Idade ≥ 60 anos, Q079 = 'Sim' e todas as outras 13 doenças = 'Não'.",
        ],
        "N": ["43.554", "4.332", "4.025", "494"],
    })
    add_df_table(doc, bases, col_widths=[3.8, 10.0, 1.7], font_size=9)
    caption(doc, "Tabela 2.2 — As quatro bases de dados derivadas da PNS 2019 e seus tamanhos amostrais.")

    para(doc,
         "Adicionalmente, sobre todas as bases foram aplicados dois filtros de integridade da "
         "entrevista: V0015 = 1 (entrevista do domicílio efetivamente realizada) e M001 = 1 "
         "(questionário individual do adulto selecionado realizado). Esses filtros garantem que "
         "apenas registros com coleta completa entrem na análise.")

    h2(doc, "2.3 Pipeline de pré-processamento")
    para(doc,
         "O pré-processamento foi estruturado em etapas sequenciais e rastreáveis, cada uma "
         "documentada em arquivos CSV de log. As principais etapas foram:")
    etapas = pd.DataFrame({
        "Etapa": ["Resolução de skip patterns", "Auditoria de missing", "Remoção de outliers",
                  "Imputação", "Fusão de atributos", "Categorização ordinal", "Encoding"],
        "Descrição": [
            "Preenchimento de valores ausentes estruturais — campos deixados em branco por "
            "roteamento condicional do questionário (ex.: quem não pratica exercício não "
            "responde 'quantos dias por semana') recebem o valor logicamente implícito.",
            "Variáveis com mais de 75% de valores ausentes reais são descartadas.",
            "Valores além de 3×IQR (calculado por classe) são substituídos por NaN.",
            "Valores ausentes reais são imputados pela média (numéricas) ou moda (categóricas) "
            "da própria classe.",
            "Criação de atributos compostos: IMC, Escore Inflamatório, Escore Saudável e a "
            "razão entre eles.",
            "Conversão de variáveis contínuas em faixas clínicas (faixa etária, IMC-OMS, nível "
            "de atividade física).",
            "Label Encoding para variáveis ordinais e One-Hot Encoding para nominais.",
        ],
    })
    add_df_table(doc, etapas, col_widths=[3.5, 12.0], font_size=8.5)
    caption(doc, "Tabela 2.3 — Etapas do pipeline de pré-processamento aplicado a ambos os desenhos.")

    callout(doc, "Sobre a resolução de skip patterns",
            "Esta etapa é um diferencial metodológico do estudo. O questionário da PNS usa "
            "roteamento condicional: a depender de uma resposta, perguntas seguintes são "
            "puladas. O campo fica vazio, mas esse vazio NÃO é dado faltante — é um valor "
            "estruturalmente determinado. Tratá-lo como missing aleatório e imputá-lo pela "
            "média introduz viés. A etapa de skip patterns preenche esses campos com o valor "
            "logicamente correto antes da auditoria de missing.", cor="D5E8F0")

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ===========================================================================
    # 3. RESULTADOS — ANÁLISE EXPLORATÓRIA E BIVARIADA
    # ===========================================================================
    h1(doc, "3. Resultados — Análise Exploratória e Bivariada")
    para(doc,
         "Esta seção apresenta a análise exploratória e bivariada comparando idosos com artrite "
         "pura contra idosos saudáveis, além da comparação descritiva entre as quatro bases. "
         "Todos os testes adotaram nível de significância α = 0,05; a notação de significância "
         "segue: *** p<0,001; ** p<0,01; * p<0,05; ns = não significativo.")

    # --- 3.1 Prevalência ---
    h2(doc, "3.1 Prevalência de artrite por subgrupo")
    tab3b = load_csv(RES_EDA / "tabelas" / "tabela_3B_prevalencia_subgrupos.csv")
    para(doc,
         "A prevalência global de artrite/reumatismo entre os idosos da PNS 2019 foi de 17,71% "
         "(IC95% 17,21–18,21). A análise por subgrupo revela um gradiente marcante por sexo e "
         "idade, conforme a Tabela 3.1 e a Figura 3.1.")
    add_df_table(doc, tab3b, font_size=8.5)
    caption(doc, "Tabela 3.1 — Prevalência de artrite por subgrupo, com intervalo de confiança de 95%.")
    add_figure(doc, RES_EDA / "figuras" / "grafico_3B_prevalencia_subgrupos.png",
               "Figura 3.1 — Prevalência de artrite por subgrupo com IC95%. A linha tracejada "
               "marca a prevalência global de 17,7%.")
    callout(doc, "Leitura dos achados — prevalência",
            "O achado mais forte é o gradiente por sexo: a prevalência em mulheres (25,1%) é "
            "quase três vezes a dos homens (8,6%). Há também gradiente etário positivo "
            "(16,6% aos 60–69 anos sobe para 20,5% aos 80+). Já a escolaridade mostra "
            "prevalência relativamente estável (15,9%–18,9%), sugerindo que o nível educacional "
            "isoladamente não é um forte marcador de risco nesta amostra. Esses padrões são "
            "consistentes com a literatura epidemiológica sobre doenças osteoarticulares e "
            "devem ser explorados na Discussão.")

    # --- 3.2 Comparação das 4 bases ---
    h2(doc, "3.2 Comparação descritiva das quatro bases")
    tab3a = load_csv(RES_EDA / "tabelas" / "tabela_3A_comparacao_quatro_bases.csv")
    para(doc,
         "A Tabela 3.2 compara as quatro bases quanto a idade, IMC, prática de exercício, "
         "consumo de ultraprocessados e proporção de mulheres. O teste de Kruskal-Wallis foi "
         "usado para as variáveis quantitativas e o qui-quadrado para as proporções. A Figura "
         "3.2 ilustra as distribuições de idade e IMC.")
    add_df_table(doc, tab3a, font_size=7.5)
    caption(doc, "Tabela 3.2 — Comparação das quatro bases (Kruskal-Wallis para quantitativas; "
                 "χ² + Odds Ratio para a proporção de mulheres).")
    add_figure(doc, RES_EDA / "figuras" / "grafico_3A_boxplots_quatro_bases.png",
               "Figura 3.2 — Distribuição de idade e IMC nas quatro bases.")
    callout(doc, "Leitura dos achados — comparação das bases",
            "O IMC é a variável que mais separa os grupos: a base 'Artrite com comorbidades' "
            "tem mediana de IMC de 26,8 kg/m², contra 24,8 da base Saudáveis (p<0,001). "
            "Interessante notar que a base 'Artrite Pura' tem IMC mediano de 25,3 — mais baixo "
            "que o da artrite com comorbidades, o que sugere que parte do excesso de peso "
            "observado em artríticos é mediado pelas comorbidades associadas (hipertensão, "
            "diabetes). A proporção feminina também é fortemente discriminante: 78,2% na "
            "artrite com comorbidades (OR 2,92 vs. base geral) e 65,0% na artrite pura "
            "(OR 1,51).")

    # --- 3.3 Variáveis quantitativas ---
    h2(doc, "3.3 Características quantitativas — Artrite Pura vs. Saudáveis")
    tab2a = load_csv(RES_EDA / "tabelas" / "tabela_2A_univariada_quantitativas.csv")
    para(doc,
         "A Tabela 3.3 apresenta as variáveis quantitativas no formato 'Tabela 1' de artigo "
         "científico (wide), com mediana, IQR, média ± DP e p-valor do teste de Mann-Whitney U "
         "para cada grupo. A Figura 3.3 traz os boxplots correspondentes.")
    add_df_table(doc, tab2a, font_size=7)
    caption(doc, "Tabela 3.3 — Características quantitativas: idosos com artrite pura vs. saudáveis.")
    add_figure(doc, RES_EDA / "figuras" / "grafico_2A_boxplots_quantitativas.png",
               "Figura 3.3 — Boxplots das variáveis quantitativas por grupo (artrite pura vs. saudáveis).")
    callout(doc, "Leitura dos achados — variáveis quantitativas",
            "Apenas três variáveis atingiram significância estatística: IMC (p<0,01, maior nos "
            "artríticos), consumo de frutas (p<0,01, maior nos artríticos) e consumo de "
            "refrigerante (p<0,05, menor nos artríticos). Os dois últimos achados são "
            "contraintuitivos sob a hipótese inflamatória — esperar-se-ia que artríticos "
            "consumissem MENOS frutas e MAIS refrigerante. A explicação mais provável é o viés "
            "de causalidade reversa: pacientes que recebem o diagnóstico modificam hábitos "
            "alimentares buscando melhora clínica. Este é um ponto central a ser discutido no "
            "artigo (ver Seção 5).")

    # --- 3.4 Variáveis qualitativas ---
    h2(doc, "3.4 Características qualitativas — Artrite Pura vs. Saudáveis")
    tab2b = load_csv(RES_EDA / "tabelas" / "tabela_2B_univariada_qualitativas.csv")
    para(doc,
         "A Tabela 3.4 detalha as variáveis categóricas, com contagem e percentual por grupo e "
         "p-valor do teste qui-quadrado. A Figura 3.4 mostra a distribuição percentual.")
    add_df_table(doc, tab2b, font_size=7.5)
    caption(doc, "Tabela 3.4 — Características qualitativas: idosos com artrite pura vs. saudáveis.")
    add_figure(doc, RES_EDA / "figuras" / "grafico_2B_barras_qualitativas.png",
               "Figura 3.4 — Distribuição percentual das variáveis qualitativas por grupo.")
    callout(doc, "Leitura dos achados — variáveis qualitativas",
            "Três variáveis foram altamente significativas (p<0,001): sexo, autoavaliação de "
            "saúde e — em menor grau — posse de plano de saúde (p<0,01). Entre os artríticos, "
            "65% são mulheres (vs. 42,3% nos saudáveis) e a moda da autoavaliação de saúde é "
            "'Regular' (vs. 'Boa' nos saudáveis). A escolaridade e a prática de exercício não "
            "apresentaram diferença significativa entre os grupos. O perfil que emerge é o de "
            "uma idosa que se percebe com saúde pior — coerente com o impacto funcional da "
            "artrite na qualidade de vida.")

    # --- 3.5 Bivariada completa ---
    h2(doc, "3.5 Análise bivariada completa")
    tab2c = load_csv(RES_EDA / "tabelas" / "tabela_2C_bivariada_completa.csv")
    para(doc,
         "A Tabela 3.5 reúne todas as variáveis (quantitativas e qualitativas) com o teste "
         "estatístico apropriado, a estatística obtida, o p-valor e o intervalo de confiança "
         "de 95% da diferença de medianas (para quantitativas, via bootstrap percentílico com "
         "2.000 réplicas).")
    add_df_table(doc, tab2c, font_size=6.5)
    caption(doc, "Tabela 3.5 — Análise bivariada completa: artrite pura vs. saudáveis.")
    callout(doc, "Síntese da análise bivariada",
            "Das 18 variáveis testadas, 6 foram estatisticamente significativas: sexo (***), "
            "autoavaliação de saúde (***), plano de saúde (**), IMC (**), consumo de frutas "
            "(**) e consumo de refrigerante (*). Notavelmente, NENHUMA variável de hábito "
            "alimentar pró-inflamatório (carne vermelha, doces, embutidos) atingiu "
            "significância — um achado que vai contra a hipótese nutricional inicial e merece "
            "discussão cuidadosa. A correlação entre as preditoras quantitativas (Tabela 2-D do "
            "notebook) não identificou nenhum par com |ρ| > 0,50, ou seja, não há "
            "multicolinearidade relevante entre as variáveis alimentares.")

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ===========================================================================
    # 4. RESULTADOS — PRÉ-PROCESSAMENTO
    # ===========================================================================
    h1(doc, "4. Resultados — Pré-Processamento e os Dois Desenhos de Estudo")

    e9_puro  = rel_puro["rastreabilidade"]["etapa_9_dataset_final"]
    e9_comor = rel_comor["rastreabilidade"]["etapa_9_dataset_final"]

    para(doc,
         "O pré-processamento foi executado em dois desenhos paralelos. A Tabela 4.1 resume os "
         "dois datasets finais lado a lado.")
    comp = pd.DataFrame({
        "Característica": ["Banco de origem", "N total", "Casos (artrite)",
                           "Controles (saudáveis)", "Razão de desbalanceamento",
                           "Nº de features finais", "Variáveis Q* como preditoras"],
        "Desenho A — Artrite Pura": [
            "idosos_artrite_puro.db",
            f"{e9_puro['n_registros']:,}".replace(",", "."),
            f"{e9_puro['distribuicao_label']['1']:,}".replace(",", "."),
            f"{e9_puro['distribuicao_label']['0']:,}".replace(",", "."),
            f"{e9_puro['razao_desbalanceamento']}:1 (severo)",
            str(e9_puro["n_features"]),
            "Não (constantes — removidas)",
        ],
        "Desenho B — Artrite c/ Comorbidades": [
            "idosos_artrite.db",
            f"{e9_comor['n_registros']:,}".replace(",", "."),
            f"{e9_comor['distribuicao_label']['1']:,}".replace(",", "."),
            f"{e9_comor['distribuicao_label']['0']:,}".replace(",", "."),
            f"{e9_comor['razao_desbalanceamento']}:1 (quase 1:1)",
            str(e9_comor["n_features"]),
            "Sim (variam — mantidas)",
        ],
    })
    add_df_table(doc, comp, col_widths=[4.0, 5.7, 5.8], font_size=8.5)
    caption(doc, "Tabela 4.1 — Comparação dos dois datasets finais produzidos pelo pré-processamento.")

    callout(doc, "Por que dois desenhos? O trade-off central",
            "O Desenho A (artrite pura) oferece um contraste limpo — compara artrite isolada "
            "contra saúde plena —, mas sofre desbalanceamento severo (8,77:1) e representa um "
            "perfil clinicamente raro (idoso com artrite e nenhuma outra doença crônica). O "
            "Desenho B (artrite com comorbidades) é quase perfeitamente balanceado (1,08:1) e "
            "clinicamente realista, mas introduz um risco de vazamento metodológico: como as "
            "demais doenças do módulo Q foram usadas como filtro para definir o grupo "
            "Saudável, elas tornam-se preditores artificialmente fortes. Apresentar os dois "
            "desenhos ao banco examinador e discutir esse trade-off é, em si, uma contribuição "
            "metodológica do trabalho.", cor="D5E8F0")

    # --- 4.1 Skip patterns ---
    h2(doc, "4.1 Resolução de skip patterns")
    sk_puro  = rel_puro["rastreabilidade"]["etapa_2_5_skip_patterns"]["total_preenchidos"]
    sk_comor = rel_comor["rastreabilidade"]["etapa_2_5_skip_patterns"]["total_preenchidos"]
    para(doc,
         f"A resolução de skip patterns preencheu {sk_puro:,} valores estruturais no Desenho A "
         f"e {sk_comor:,} no Desenho B".replace(",", ".") +
         ". Em ambos os casos, o grosso da correção concentra-se na variável J012 (número de "
         "consultas médicas nos últimos 12 meses): idosos que declararam não ter consultado "
         "médico no último ano tinham esse campo vazio, agora corretamente preenchido com 0.")
    tab_skip = load_csv(RES_COMOR / "tabelas" / "etapa2_5_skip_patterns.csv")
    add_df_table(doc, tab_skip, font_size=8)
    caption(doc, "Tabela 4.2 — Mapa de skip patterns resolvidos (valores do Desenho B).")

    # --- 4.2 Missing + outliers + imputação ---
    h2(doc, "4.2 Auditoria de missing, outliers e imputação")
    ve_puro  = rel_puro["rastreabilidade"]["etapa_3_missing"]["n_excluidas"]
    ve_comor = rel_comor["rastreabilidade"]["etapa_3_missing"]["n_excluidas"]
    out_puro = rel_puro["rastreabilidade"]["etapa_4_outliers"]["total_outliers_substituidos"]
    out_comor = rel_comor["rastreabilidade"]["etapa_4_outliers"]["total_outliers_substituidos"]
    imp_puro = rel_puro["rastreabilidade"]["etapa_5_imputacao"]["total_imputados"]
    imp_comor = rel_comor["rastreabilidade"]["etapa_5_imputacao"]["total_imputados"]
    para(doc,
         f"No Desenho A foram excluídas {ve_puro} variáveis por excesso de missing (>75%), "
         f"{out_puro} outliers substituídos por NaN e {imp_puro:,} valores imputados. ".replace(",", ".") +
         f"No Desenho B, {ve_comor} variáveis excluídas, {out_comor} outliers e "
         f"{imp_comor:,} valores imputados. ".replace(",", ".") +
         "A Figura 4.1 mostra o mapa de calor de dados ausentes por variável e grupo (Desenho B).")
    add_figure(doc, RES_COMOR / "figuras" / "etapa3_heatmap_missing.png",
               "Figura 4.1 — Mapa de calor de dados ausentes reais por variável e grupo (Desenho B), "
               "após a resolução de skip patterns.", largura_cm=12.0)
    add_figure(doc, RES_COMOR / "figuras" / "etapa4_boxplots_outliers.png",
               "Figura 4.2 — Boxplots antes e depois da remoção de outliers (Desenho B).")

    # --- 4.3 Análise de comorbidades ---
    h2(doc, "4.3 Caracterização da multimorbidade nos idosos com artrite")
    tab_comorb = load_csv(RES_COMOR / "tabelas" / "etapa3_5_comorbidades_prevalencia.csv")
    para(doc,
         "Uma análise exclusiva do Desenho B caracteriza quantas e quais comorbidades os idosos "
         "com artrite apresentam. Este é um dos achados descritivos mais relevantes para a "
         "Discussão do artigo: o idoso com artrite é, tipicamente, um idoso multimórbido.")
    add_df_table(doc, tab_comorb, col_widths=[2.5, 5.0, 2.0, 2.0], font_size=9)
    caption(doc, "Tabela 4.3 — Prevalência de cada comorbidade entre os 4.025 idosos com artrite.")
    add_figure(doc, RES_COMOR / "figuras" / "etapa3_5_comorbidades.png",
               "Figura 4.3 — Distribuição do número de comorbidades por idoso com artrite (esquerda) "
               "e prevalência de cada doença crônica associada (direita).")
    callout(doc, "Leitura dos achados — multimorbidade",
            "Entre os 4.025 idosos com artrite, 65,3% também têm hipertensão, 39,8% colesterol "
            "alto, 21,9% diabetes e 19,5% depressão. Apenas uma fração pequena tem artrite "
            "isolada (os 494 do Desenho A). Isso fundamenta um argumento importante para o "
            "artigo: estudar a 'artrite pura' representa um perfil clinicamente atípico — na "
            "prática, o idoso brasileiro com artrite carrega, em média, mais de uma "
            "comorbidade. A associação artrite–depressão (19,5%) é especialmente relevante por "
            "dialogar diretamente com Cancella & Zárate (2025) e Xue et al. (2020).")

    # --- 4.4 Diagnóstico dos datasets finais ---
    h2(doc, "4.4 Diagnóstico dos datasets finais")
    para(doc,
         "As Figuras 4.4 e 4.5 mostram o diagnóstico final de cada dataset: distribuição das "
         "classes, IMC por grupo e escores alimentares por grupo.")
    add_figure(doc, RES_PURO / "figuras" / "etapa9_diagnostico_dataset_final.png",
               "Figura 4.4 — Diagnóstico do dataset final do Desenho A (artrite pura). Note o "
               "desbalanceamento severo entre as classes.")
    add_figure(doc, RES_COMOR / "figuras" / "etapa9_diagnostico_dataset_final.png",
               "Figura 4.5 — Diagnóstico do dataset final do Desenho B (artrite com comorbidades). "
               "As classes estão quase perfeitamente balanceadas.")

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ===========================================================================
    # 5. DISCUSSÕES PROPOSTAS
    # ===========================================================================
    h1(doc, "5. Discussões Propostas para o Artigo")
    para(doc,
         "Esta seção reúne os pontos de discussão mais férteis identificados a partir dos "
         "resultados. Cada item é uma semente de parágrafo para a seção de Discussão do artigo.")

    h2(doc, "5.1 O paradoxo nutricional e a causalidade reversa")
    para(doc,
         "A hipótese inicial do estudo previa que idosos com artrite teriam um padrão alimentar "
         "mais pró-inflamatório. Os dados, porém, mostram o contrário em duas variáveis "
         "significativas: artríticos consomem MAIS frutas e MENOS refrigerante que os "
         "saudáveis. A explicação mais plausível é a causalidade reversa, inerente ao "
         "delineamento transversal: o diagnóstico de artrite leva o paciente a modificar a "
         "dieta. Esta limitação deve ser explicitamente discutida e ancorada na literatura "
         "(Hernán et al., 2004, sobre viés de seleção e causalidade reversa).")

    h2(doc, "5.2 O peso (IMC) como elo entre artrite e comorbidades")
    para(doc,
         "O IMC foi significativo em todas as análises. Mais interessante é o gradiente entre "
         "as bases: artrite pura (25,3) < artrite com comorbidades (26,8). Isso sugere que o "
         "excesso de peso observado em artríticos é, em parte, mediado pelas comorbidades "
         "metabólicas associadas (hipertensão, diabetes). Uma análise de mediação — ou ao menos "
         "uma discussão qualitativa desse mecanismo — enriqueceria o artigo.")

    h2(doc, "5.3 O perfil sociodemográfico: a artrite tem gênero e idade")
    para(doc,
         "O achado mais robusto de todo o estudo é o gradiente por sexo: mulheres têm "
         "prevalência de artrite quase três vezes maior que homens (25,1% vs. 8,6%). Somado ao "
         "gradiente etário positivo, o perfil de risco que emerge é claro e consistente com a "
         "literatura internacional. Esse achado, por si só, valida a robustez da base e do "
         "pipeline de extração.")

    h2(doc, "5.4 O dilema metodológico dos dois desenhos")
    para(doc,
         "A decisão de construir dois datasets — artrite pura e artrite com comorbidades — "
         "permite uma discussão metodológica madura. O Desenho A oferece pureza de contraste; "
         "o Desenho B oferece realismo clínico e balanceamento. Recomenda-se, no Notebook de "
         "modelagem, treinar os algoritmos em ambos e, para o Desenho B, conduzir uma análise "
         "de sensibilidade: um Modelo com as variáveis Q* e um Modelo sem elas, para quantificar "
         "quanto da acurácia advém do vazamento de delineamento.")

    h2(doc, "5.5 A multimorbidade como achado clínico central")
    para(doc,
         "A constatação de que 2 em cada 3 idosos com artrite também têm hipertensão — e que a "
         "artrite isolada é rara — reposiciona o estudo. Mais do que classificar 'artrite vs. "
         "saúde', o trabalho descreve um perfil de fragilidade multidimensional. A associação "
         "com depressão (19,5%) conecta-se diretamente à linha de pesquisa do orientador "
         "(Cancella & Zárate, 2025).")

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ===========================================================================
    # 6. LIMITAÇÕES
    # ===========================================================================
    h1(doc, "6. Limitações Metodológicas")
    for titulo, txt in [
        ("Delineamento transversal",
         "Exposição e desfecho são medidos simultaneamente, impossibilitando inferência causal "
         "e tornando o estudo vulnerável à causalidade reversa (ver Seção 5.1)."),
        ("Dados autorreferidos",
         "O diagnóstico de artrite e os hábitos alimentares são autorreferidos, sujeitos a "
         "viés de memória e de desejabilidade social. A PNS não dispõe de biomarcadores "
         "inflamatórios laboratoriais."),
        ("Vazamento de delineamento no Desenho B",
         "As variáveis do módulo Q foram usadas como filtro para definir o grupo Saudável, o "
         "que as torna preditores artificialmente fortes no Desenho B. Mitigação proposta: "
         "análise de sensibilidade com e sem essas variáveis."),
        ("Imputação por classe e risco de leakage para ML",
         "A imputação usou estatísticas da própria classe, o que é adequado para a análise "
         "descritiva mas vaza informação do alvo para as features. Na modelagem, a imputação "
         "deve ser refeita dentro de cada fold de validação cruzada."),
        ("Desbalanceamento no Desenho A",
         "A razão de 8,77:1 exige balanceamento artificial (Random Under Sampler), que reduz o "
         "tamanho efetivo da amostra de treino e pode aumentar a variância das métricas."),
        ("Variáveis descartadas por missing",
         "Variáveis potencialmente relevantes (renda — VDF004; autoavaliação detalhada — N001; "
         "escolaridade — VDD004A) foram descartadas por terem 100% de missing real no recorte "
         "analítico, limitando a dimensão socioeconômica da análise."),
    ]:
        bullet(doc, txt, bold_prefix=titulo + ".")

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ===========================================================================
    # 7. PRÓXIMOS PASSOS
    # ===========================================================================
    h1(doc, "7. Próximos Passos — Modelagem de Machine Learning")
    para(doc,
         "Com os dois datasets prontos, a próxima etapa é a fase de modelagem do processo KDD. "
         "O plano de execução, alinhado com a metodologia da linha de pesquisa, é:")
    for txt in [
        "Treinar quatro algoritmos supervisionados em cada desenho: Árvore de Decisão e Naive "
        "Bayes (caixa-branca, interpretáveis) e Floresta Aleatória e AdaBoost (ensemble, maior "
        "desempenho).",
        "Aplicar Random Under Sampler dentro de cada fold da validação cruzada (especialmente "
        "crítico no Desenho A).",
        "Otimizar hiperparâmetros com RandomizedSearchCV e validação cruzada estratificada de "
        "10 folds; reportar F1-score com IC 95% por classe.",
        "Avaliar no conjunto de teste (holdout) com acurácia, AUC-ROC e matriz de confusão.",
        "Comparar estatisticamente os algoritmos com teste-T sobre os F1-scores dos 10 folds.",
        "Extrair Feature Importance e regras da Árvore de Decisão (profundidade limitada) para "
        "interpretabilidade.",
        "Para o Desenho B, conduzir a análise de sensibilidade Modelo A (com Q*) vs. Modelo B "
        "(sem Q*) descrita na Seção 5.4.",
    ]:
        bullet(doc, txt)

    callout(doc, "Recomendação para a apresentação ao orientador",
            "Apresentar os dois desenhos como uma escolha metodológica deliberada, não como "
            "indecisão. O Desenho A responde 'o que distingue a artrite isolada?'; o Desenho B "
            "responde 'como diferenciar o idoso multimórbido com artrite do idoso plenamente "
            "saudável?'. São perguntas científicas distintas e complementares.", cor="D5E8F0")

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ===========================================================================
    # APÊNDICE
    # ===========================================================================
    h1(doc, "Apêndice A — Índice de Tabelas e Figuras")
    h2(doc, "Tabelas")
    for t in [
        "Tabela 2.1 — Estrutura PICOS da pergunta de pesquisa.",
        "Tabela 2.2 — As quatro bases de dados derivadas da PNS 2019.",
        "Tabela 2.3 — Etapas do pipeline de pré-processamento.",
        "Tabela 3.1 — Prevalência de artrite por subgrupo com IC95%.",
        "Tabela 3.2 — Comparação das quatro bases.",
        "Tabela 3.3 — Características quantitativas: artrite pura vs. saudáveis.",
        "Tabela 3.4 — Características qualitativas: artrite pura vs. saudáveis.",
        "Tabela 3.5 — Análise bivariada completa.",
        "Tabela 4.1 — Comparação dos dois datasets finais.",
        "Tabela 4.2 — Mapa de skip patterns resolvidos.",
        "Tabela 4.3 — Prevalência de comorbidades entre idosos com artrite.",
    ]:
        p = doc.add_paragraph(t)
        p.paragraph_format.space_after = Pt(2)
    h2(doc, "Figuras")
    for f in [
        "Figura 3.1 — Prevalência de artrite por subgrupo com IC95%.",
        "Figura 3.2 — Distribuição de idade e IMC nas quatro bases.",
        "Figura 3.3 — Boxplots das variáveis quantitativas por grupo.",
        "Figura 3.4 — Distribuição percentual das variáveis qualitativas.",
        "Figura 4.1 — Mapa de calor de dados ausentes (Desenho B).",
        "Figura 4.2 — Boxplots antes/depois da remoção de outliers (Desenho B).",
        "Figura 4.3 — Distribuição e prevalência de comorbidades.",
        "Figura 4.4 — Diagnóstico do dataset final do Desenho A.",
        "Figura 4.5 — Diagnóstico do dataset final do Desenho B.",
    ]:
        p = doc.add_paragraph(f)
        p.paragraph_format.space_after = Pt(2)

    # Rodapé final
    doc.add_paragraph()
    fim = doc.add_paragraph()
    fim.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = fim.add_run("— Fim do documento consolidado de resultados —")
    rf.italic = True
    rf.font.size = Pt(9)
    rf.font.color.rgb = CINZA

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(f"[OK] Documento gerado: {OUT_PATH}")
    print(f"     {len(doc.paragraphs)} paragrafos, {len(doc.tables)} tabelas")


if __name__ == "__main__":
    build()
