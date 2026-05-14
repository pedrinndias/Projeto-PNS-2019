# -*- coding: utf-8 -*-
"""Gera o template de trabalho do artigo: estrutura do modelo do Prof. Zarate
com figuras e tabelas REAIS posicionadas, e em cada secao uma caixa de
instrucao ("O QUE ESCREVER AQUI") derivada do guia de redacao.

NAO contem a prosa do artigo. Contem o andaime + o material, para o autor
escrever cada secao com a propria voz.

Uso (da raiz do projeto):  python scripts/build_template_trabalho.py
Gera: Documentos_organizacao/Artigo_Template_de_Trabalho_PNS2019.docx
"""
from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
RES_EDA   = ROOT / "data" / "results" / "eda"
RES_PURO  = ROOT / "data" / "results" / "preprocessing"
RES_COMOR = ROOT / "data" / "results" / "preprocessing_comorbidades"
OUT_PATH  = ROOT / "Documentos_organizacao" / "Artigo_Template_de_Trabalho_PNS2019.docx"

AZUL_ESCURO = RGBColor(0x1F, 0x37, 0x64)
AZUL_MEDIO  = RGBColor(0x2E, 0x75, 0xB6)
CINZA       = RGBColor(0x59, 0x59, 0x59)
PRETO       = RGBColor(0x00, 0x00, 0x00)

# cores de caixa
C_INSTRUCAO = "E8F0FE"   # azul claro — o que escrever
C_PENDENTE  = "FCE8E6"   # vermelho claro — depende do notebook 04
C_CUIDADO   = "FEF3CD"   # amarelo — armadilha
C_PRONTO    = "E6F4EA"   # verde — material pronto


# ───────────────────────── helpers ─────────────────────────
def setup_styles(doc):
    n = doc.styles["Normal"]
    n.font.name = "Calibri"; n.font.size = Pt(11)
    n._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    for lvl, size in [(1, 16), (2, 13), (3, 11.5)]:
        st = doc.styles[f"Heading {lvl}"]
        st.font.name = "Calibri"; st.font.size = Pt(size); st.font.bold = True
        st.font.color.rgb = AZUL_ESCURO if lvl == 1 else AZUL_MEDIO


def h1(doc, t): return doc.add_heading(t, level=1)
def h2(doc, t): return doc.add_heading(t, level=2)
def h3(doc, t): return doc.add_heading(t, level=3)


def para(doc, txt, italic=False, bold=False, size=11, justify=True, color=None):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(txt)
    r.italic = italic; r.bold = bold; r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p


def caption(doc, txt):
    p = doc.add_paragraph()
    r = p.add_run(txt)
    r.italic = True; r.font.size = Pt(8.5); r.font.color.rgb = CINZA
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def _shade(cell, hex_color):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): hex_color})
    tcpr.append(shd)


def box(doc, titulo, blocos, cor=C_INSTRUCAO):
    """Caixa de instrucao. 'blocos' = lista de (tipo, conteudo).
    tipos: 'sub' (subtitulo bold), 'txt' (paragrafo), 'q' (pergunta norteadora),
           'b' (bullet)."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    _shade(cell, cor)
    cell.text = ""
    p0 = cell.paragraphs[0]
    r0 = p0.add_run(titulo)
    r0.bold = True; r0.font.size = Pt(10.5); r0.font.color.rgb = AZUL_ESCURO
    for tipo, conteudo in blocos:
        p = cell.add_paragraph()
        if tipo == "sub":
            r = p.add_run(conteudo); r.bold = True; r.font.size = Pt(9.5)
        elif tipo == "txt":
            r = p.add_run(conteudo); r.font.size = Pt(9.5)
        elif tipo == "q":
            rb = p.add_run("?  "); rb.bold = True; rb.font.color.rgb = AZUL_MEDIO
            rb.font.size = Pt(9.5)
            r = p.add_run(conteudo); r.font.size = Pt(9.5)
            p.paragraph_format.left_indent = Cm(0.3)
        elif tipo == "b":
            rb = p.add_run("•  "); rb.font.size = Pt(9.5)
            r = p.add_run(conteudo); r.font.size = Pt(9.5)
            p.paragraph_format.left_indent = Cm(0.3)
    doc.add_paragraph()


def write_here(doc):
    """Marcador de onde o autor escreve."""
    p = doc.add_paragraph()
    r = p.add_run("✍  [ESCREVA AQUI O TEXTO DESTA SEÇÃO]")
    r.italic = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    p.paragraph_format.space_after = Pt(14)


def add_figure(doc, path: Path, legenda, largura_cm=14.5):
    if not path.exists():
        para(doc, f"[FIGURA NÃO ENCONTRADA: {path.name}]", italic=True, color=CINZA)
        return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(largura_cm))
    caption(doc, legenda)


def add_df(doc, df: pd.DataFrame, legenda, font_size=7.5, col_widths=None, max_rows=None):
    if max_rows:
        df = df.head(max_rows)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr[i].text = str(col)
        for par in hdr[i].paragraphs:
            for run in par.runs:
                run.bold = True; run.font.size = Pt(font_size)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade(hdr[i], "2E75B6")
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = "" if pd.isna(val) else str(val)
            for par in cells[i].paragraphs:
                for run in par.runs:
                    run.font.size = Pt(font_size)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    caption(doc, legenda)


def add_rows_table(doc, headers, rows, legenda, font_size=8, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, txt in enumerate(headers):
        hdr[i].text = txt
        for par in hdr[i].paragraphs:
            for run in par.runs:
                run.bold = True; run.font.size = Pt(font_size)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade(hdr[i], "1F3764")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for par in cells[i].paragraphs:
                for run in par.runs:
                    run.font.size = Pt(font_size)
    if col_widths:
        for r in table.rows:
            for i, w in enumerate(col_widths):
                r.cells[i].width = Cm(w)
    caption(doc, legenda)


def load_csv(path):
    df = pd.read_csv(path)
    if df.columns[0].startswith("Unnamed"):
        df = df.drop(columns=df.columns[0])
    return df


def pagebreak(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# ───────────────────────── build ─────────────────────────
def build():
    doc = Document()
    setup_styles(doc)
    for s in doc.sections:
        s.top_margin = Cm(2.2); s.bottom_margin = Cm(2.2)
        s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

    rel_puro  = json.loads((RES_PURO / "relatorio_preprocessamento.json").read_text(encoding="utf-8"))
    rel_comor = json.loads((RES_COMOR / "relatorio_preprocessamento.json").read_text(encoding="utf-8"))

    # ===== CABEÇALHO DO ARTIGO =====
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Padrões Nutricionais e sua Associação com Artrite e Reumatismo em "
                  "Idosos Brasileiros: Uma Abordagem por Aprendizado de Máquina com "
                  "Dados da Pesquisa Nacional de Saúde 2019")
    r.bold = True; r.font.size = Pt(14); r.font.color.rgb = PRETO
    te = doc.add_paragraph(); te.alignment = WD_ALIGN_PARAGRAPH.CENTER
    re = te.add_run("Dietary Patterns and Their Association with Arthritis and Rheumatism "
                    "in Brazilian Older Adults: A Machine Learning Approach Using Data "
                    "from the 2019 National Health Survey")
    re.italic = True; re.font.size = Pt(11); re.font.color.rgb = CINZA
    doc.add_paragraph()
    a = doc.add_paragraph(); a.alignment = WD_ALIGN_PARAGRAPH.CENTER
    a.add_run("Pedro Dias Soares, Luis Enrique Zárate Gálvez").font.size = Pt(11)
    af = doc.add_paragraph(); af.alignment = WD_ALIGN_PARAGRAPH.CENTER
    af.add_run("Curso de Ciência de Dados — Pontifícia Universidade Católica de Minas "
               "Gerais (PUC Minas)\npdsoares@sga.pucminas.br, zarate@pucminas.br").font.size = Pt(9.5)
    doc.add_paragraph()

    box(doc, "📄 SOBRE ESTE DOCUMENTO", [
        ("txt", "Este é o TEMPLATE DE TRABALHO do seu artigo. Ele tem a estrutura exata do "
                "modelo do Prof. Zárate, com as figuras e tabelas reais já posicionadas e, em "
                "cada seção, uma caixa azul «O QUE ESCREVER AQUI» com instruções, perguntas "
                "norteadoras e o material disponível."),
        ("txt", "A prosa é sua. As caixas azuis e o marcador ✍ indicam onde e o que você "
                "escreve. Caixas vermelhas = pendência (depende do notebook 04 de ML). "
                "Caixas amarelas = armadilha/atenção."),
        ("txt", "Ordem recomendada de escrita: Métodos → Resultados → Discussão/Conclusão → "
                "Introdução → Resumo. Não escreva na ordem de leitura."),
    ], cor=C_INSTRUCAO)
    pagebreak(doc)

    # ===== RESUMO / ABSTRACT =====
    h1(doc, "Resumo")
    box(doc, "📝 O QUE ESCREVER AQUI — Resumo (escrever POR ÚLTIMO)", [
        ("sub", "Objetivo da seção:"),
        ("txt", "Miniatura informativa do artigo: Objetivo + Método + Resultados + Conclusão, "
                "em 150–250 palavras. O JHI exige versões em PT, EN e ES."),
        ("sub", "Perguntas norteadoras:"),
        ("q", "Consegue dizer o objetivo em 1–2 frases? E o método (KDD, CAPTO, algoritmos, "
              "PNS 2019)? E o resultado principal? E a conclusão?"),
        ("q", "O objetivo e a conclusão do resumo combinam exatamente entre si?"),
        ("sub", "Por que por último:"),
        ("txt", "É impossível resumir fielmente um artigo que ainda não está pronto. Deixe "
                "esta seção em branco até o fim."),
    ])
    write_here(doc)
    h2(doc, "Descritores")
    para(doc, "Sugestão (confirmar no DeCS): Artrite; Mineração de Dados; Aprendizado de "
              "Máquina; Saúde do Idoso; Pesquisa Nacional de Saúde.", italic=True, color=CINZA)
    h1(doc, "Abstract")
    box(doc, "📝 WRITE HERE — Abstract", [
        ("txt", "Tradução fiel do Resumo em português. Mesma estrutura: Objective + Method + "
                "Results + Conclusion."),
    ])
    write_here(doc)
    pagebreak(doc)

    # ===== 1. INTRODUÇÃO =====
    h1(doc, "1. Introdução")
    box(doc, "📝 O QUE ESCREVER AQUI — Introdução (escrever em 5º lugar)", [
        ("sub", "O que o gabarito exige (itens 1–4):"),
        ("b", "Item 1: domínio contextualizado com dados epidemiológicos de impacto."),
        ("b", "Item 2: trabalhos relacionados citados brevemente (não a seção completa)."),
        ("b", "Item 3: problema, fonte de dados (PNS 2019) e técnica mencionados."),
        ("b", "Item 4: último parágrafo descreve a estrutura do artigo por seções."),
        ("sub", "Material que você já tem:"),
        ("b", "Um dado de impacto CALCULADO POR VOCÊ: prevalência de 17,71% de artrite entre "
              "idosos (mais forte que citar literatura — é seu resultado)."),
        ("b", "Contexto e estado da arte: Plano de Artigo (problema, estado da arte, lacuna)."),
        ("b", "Objetivo já redigido no Plano de Artigo."),
        ("sub", "REGRA DE OURO:"),
        ("txt", "O objetivo fica no FINAL da Introdução, e deve combinar EXATAMENTE com a "
                "Conclusão."),
    ])
    h2(doc, "1.1 Contextualização e Impacto")
    box(doc, "📝 O QUE ESCREVER AQUI", [
        ("q", "Qual dado de impacto abre o artigo? Sugestão: combine a prevalência que você "
              "calculou (17,71%) com dados da OMS/literatura sobre carga de doenças "
              "osteoarticulares."),
        ("q", "Por que a artrite em idosos brasileiros é um problema de saúde pública "
              "relevante o bastante para um artigo?"),
        ("q", "Como o envelhecimento populacional brasileiro torna o tema urgente?"),
    ])
    write_here(doc)
    h2(doc, "1.2 Trabalhos Relacionados (breve menção)")
    box(doc, "📝 O QUE ESCREVER AQUI", [
        ("q", "Em 2–3 frases, que trabalhos anteriores você menciona? (Cancella & Zárate "
              "2025; Melo et al. 2024; Silva & Zárate 2024.) Aqui é só menção — a seção "
              "completa é a 2."),
    ])
    write_here(doc)
    h2(doc, "1.3 Objetivo")
    box(doc, "📝 O QUE ESCREVER AQUI", [
        ("txt", "Use o objetivo do Plano de Artigo como base."),
        ("q", "O seu objetivo final combina EXATAMENTE com a conclusão que você escreveu na "
              "etapa anterior? Coloque os dois lado a lado e confira."),
    ])
    write_here(doc)
    h2(doc, "1.4 Estrutura do Artigo")
    para(doc, "Este trabalho está estruturado em cinco seções principais. Na Seção 2 são "
              "apresentados os trabalhos relacionados ao tema. Na Seção 3, a metodologia "
              "adotada é descrita, incluindo materiais e todas as etapas do processo de "
              "descoberta de conhecimento. Na Seção 4, os experimentos e a análise dos "
              "resultados são discutidos. A Seção 5 apresenta as conclusões e os trabalhos "
              "futuros.")
    box(doc, "📝 AJUSTE", [("txt", "O parágrafo acima veio do template. Confira se ainda bate "
                                   "com a sua estrutura real antes de finalizar.")],
        cor=C_CUIDADO)
    pagebreak(doc)

    # ===== 2. TRABALHOS RELACIONADOS =====
    h1(doc, "2. Trabalhos Relacionados")
    box(doc, "📝 O QUE ESCREVER AQUI — Trabalhos Relacionados (pode escrever junto com Métodos)", [
        ("sub", "O que o gabarito exige (itens 5–6):"),
        ("b", "Item 5: mínimo de 6 trabalhos (conferência e periódico, nacionais e internacionais)."),
        ("b", "Item 6: apontar a diferença do seu trabalho em relação à literatura."),
        ("sub", "Material que você já tem — o Plano de Artigo lista 8 referências com o motivo "
                "de usar cada uma:"),
        ("b", "Cancella & Zárate (2025) — perfil artrite-depressão, mesma base e metodologia."),
        ("b", "Melo et al. (2024) — DPOC com Random Forest, CAPTO + PNS 2019 (modelo metodológico)."),
        ("b", "Silva & Zárate (2024) — depressão, F1=82% com Floresta Aleatória."),
        ("b", "Gonçalves & Zárate (2024) — perfis de AVC, Árvore de Decisão."),
        ("b", "Zárate et al. (2023) — método CAPTO."),
        ("b", "Xue et al. (2020) — artrite aumenta 35–50% o risco de depressão."),
        ("sub", "Perguntas norteadoras:"),
        ("q", "Para cada trabalho: contexto → técnica → base de dados → resultado → conclusão. "
              "Consegue resumir cada um nesse padrão de 1–2 frases?"),
        ("q", "Qual a diferença do SEU trabalho? (Foco específico em artrite na população "
              "idosa brasileira; DOIS desenhos de estudo paralelos — pura vs. com "
              "comorbidades; método CAPTO.)"),
    ])
    h2(doc, "2.1 Posicionamento deste Trabalho")
    write_here(doc)
    pagebreak(doc)

    # ===== 3. MATERIAIS E MÉTODOS =====
    h1(doc, "3. Materiais e Métodos")
    box(doc, "📝 ORIENTAÇÃO GERAL — Materiais e Métodos (ESCREVER PRIMEIRO)", [
        ("txt", "Seção mais factual e a mais alimentada pelo trabalho dos notebooks. Não há "
                "interpretação aqui — você descreve, com precisão, o que foi feito. O Prof. "
                "Zárate exige Materiais e Métodos separados (item 7) e as 10 etapas do "
                "pipeline individualizadas (itens 11–22)."),
        ("txt", "Boa notícia: você consegue redigir AGORA tudo da Etapa 1 à Etapa 6. As "
                "Etapas 7–10 dependem do notebook 04 (ML) — estão marcadas em vermelho."),
    ])

    # 3.1 Materiais
    h2(doc, "3.1 Materiais — A Base de Dados PNS 2019")
    box(doc, "📝 O QUE ESCREVER AQUI", [
        ("sub", "Gabarito: item 9 (base descrita) + item 10 (análise descritiva)."),
        ("sub", "Perguntas norteadoras:"),
        ("q", "Como você explica em 3–4 frases o que é a PNS 2019 e por que ela serve ao seu "
              "objetivo? (Cobertura nacional, dados de saúde autorreferidos, domínio público, "
              "293.726 registros, 1.088 atributos.)"),
        ("q", "Por que você construiu QUATRO bases e não duas? O que cada uma representa "
              "clinicamente?"),
        ("q", "O que são os filtros V0015=1 e M001=1 e por que sem eles a amostra ficaria "
              "contaminada com entrevistas incompletas?"),
    ])
    add_rows_table(doc, ["Base", "Critério de construção", "N"], [
        ["Base 1 — População idosa geral", "Idade ≥ 60 anos, registros completos dos módulos usados", "43.554"],
        ["Base 2 — Idosos saudáveis", "Idade ≥ 60 e 'Não' para todas as 14 doenças crônicas do módulo Q", "4.332"],
        ["Base 3 — Artrite com comorbidades", "Idade ≥ 60 e Q079='Sim', em qualquer combinação de comorbidades", "4.025"],
        ["Base 4 — Artrite pura", "Idade ≥ 60, Q079='Sim' e todas as outras 13 doenças='Não'", "494"],
    ], "Tabela 3.1 — As quatro bases de dados derivadas da PNS 2019 (resultado dos notebooks 01 e 03/03B).",
       col_widths=[4.5, 9.0, 1.8])
    write_here(doc)

    # 3.2 Etapa 1
    h2(doc, "3.2 Etapa 1 — Entendimento do Problema")
    box(doc, "📝 O QUE ESCREVER AQUI", [
        ("txt", "Material pronto: a estrutura PICOS está no Plano de Artigo — use-a literalmente."),
        ("q", "Em uma frase: qual é o problema que justifica este estudo existir?"),
        ("q", "Qual é a pergunta de pesquisa completa? (Plano de Artigo tem a versão final.)"),
    ])
    add_rows_table(doc, ["Componente", "Definição no estudo"], [
        ["P — População", "Idosos com 60 anos ou mais registrados na PNS 2019."],
        ["I — Exposição", "Padrões nutricionais e variáveis sociodemográficas, físicas, mentais e antropométricas."],
        ["C — Comparação", "Idosos saudáveis, sem diagnóstico de artrite nem de outras doenças crônicas."],
        ["O — Desfecho", "Classificação do perfil: Saudável vs. Portador de Artrite/Reumatismo (F1-score por classe)."],
        ["S — Tipo de estudo", "Estudo transversal com processo KDD + aprendizado de máquina supervisionado."],
    ], "Tabela 3.2 — Estrutura PICOS da pergunta de pesquisa (fonte: Plano de Artigo).",
       col_widths=[3.2, 12.0])
    write_here(doc)

    # 3.3 Etapa 2 — CAPTO
    h2(doc, "3.3 Etapa 2 — Entendimento do Domínio e Modelagem Conceitual (CAPTO)")
    box(doc, "📝 O QUE ESCREVER AQUI", [
        ("sub", "Gabarito: item 13 — entendimento do domínio e modelagem conceitual (CAPTO)."),
        ("sub", "Material pronto:"),
        ("b", "A Tabela 3.3 abaixo resume o Mapa Conceitual CAPTO por dimensão (do Guia de "
              "Análises e do Plano de Artigo)."),
        ("b", "A Tabela 18 do Plano de Artigo tem a versão completa, atributo por atributo, "
              "com fonte científica e estatísticas — use-a para detalhar."),
        ("sub", "Perguntas norteadoras:"),
        ("q", "Quais são as 6 dimensões conceituais do problema? Para cada uma, qual a "
              "justificativa científica de tê-la incluído e que referência a sustenta?"),
        ("q", "Como cada dimensão se conecta ao desfecho (artrite)? É essa lógica que a "
              "Figura 1 (Modelo Conceitual) precisa mostrar visualmente."),
    ])
    add_rows_table(doc, ["Dimensão", "Aspectos", "Atributos PNS associados"], [
        ["Hábitos de saúde", "Atividade física, comportamento sedentário",
         "P034, P035, P04501"],
        ["Hábitos alimentares", "Ingestão de alimentos in natura e ultraprocessados; bebidas",
         "P00901, P015, P018, P01101, P02501, P02002, P02001, P023, P01601, P02602, P00603, P00620"],
        ["Acesso e uso de serviços de saúde", "Última consulta, frequência, plano de saúde",
         "J01101, J012, J01002, I00101"],
        ["Condições físicas e mentais", "Deficiências, saúde mental, autoavaliação, doenças "
         "crônicas, dados laboratoriais",
         "G059, G060, G062, Q092, Q11006, N001, Q00201, Q03001, Q060, Q06306, Q068, Q074, Q084, Q120, Q124, Q11604"],
        ["Dados sociodemográficos", "Sexo, renda, idade, escolaridade",
         "C006, VDF004, C008, VDD004A"],
        ["Antropometria", "Peso e altura (insumos do IMC)", "P00104, P00404"],
    ], "Tabela 3.3 — Mapa Conceitual CAPTO: dimensões, aspectos e atributos da PNS 2019.",
       col_widths=[3.3, 5.0, 6.9], font_size=7.5)
    box(doc, "⚠️ PENDÊNCIA — Figura 1 (Modelo Conceitual)", [
        ("txt", "O template pede a Figura 1 — o Mapa Conceitual visual ligando as 6 dimensões "
                "ao atributo-alvo (artrite). NÃO temos essa figura: ela é conceitual, não sai "
                "de código. Você precisa desenhá-la (PowerPoint, draw.io ou à mão) a partir "
                "da Tabela 3.3 acima e da Tabela 18 do Plano de Artigo. Insira-a aqui quando "
                "estiver pronta."),
    ], cor=C_PENDENTE)
    para(doc, "[FIGURA 1 — Inserir aqui: Modelo Conceitual CAPTO para Artrite e Reumatismo]",
         italic=True, color=CINZA, justify=False)
    caption(doc, "Figura 1 — Modelo Conceitual CAPTO para Artrite e Reumatismo em idosos.")
    write_here(doc)

    # 3.4 Etapa 3 — Seleção de atributos
    h2(doc, "3.4 Etapa 3 — Seleção Conceitual de Atributos")
    box(doc, "📝 O QUE ESCREVER AQUI", [
        ("sub", "Gabarito: item 14 — seleção conceitual de atributos COM TABELA (é a Tabela 1 "
                "do artigo)."),
        ("sub", "Perguntas norteadoras:"),
        ("q", "Quantos atributos você selecionou no total? O template pede no mínimo 10 + o "
              "alvo — você tem bem mais."),
        ("q", "Para a coluna 'Razão de Seleção': por que CADA atributo foi escolhido? O que "
              "ele captura do domínio? (Use a Tabela 18 do Plano de Artigo.)"),
        ("q", "Qual é o atributo-alvo? (Q079 — diagnóstico de artrite/reumatismo.)"),
    ])
    add_rows_table(doc, ["Código PNS", "Variável", "Tipo estatístico", "Papel no ML"], [
        ["Q079", "Diagnóstico de artrite ou reumatismo", "Qualitativa nominal binária", "Variável-alvo (Y)"],
        ["C006", "Sexo", "Qualitativa nominal", "Preditora X / controle"],
        ["C008", "Idade (anos)", "Quantitativa discreta", "Preditora X / filtro"],
        ["VDD004A", "Escolaridade", "Qualitativa ordinal", "Preditora X / controle"],
        ["VDF004", "Faixa de renda per capita", "Quantitativa contínua", "Preditora X / controle"],
        ["P00104 / P00404", "Peso e altura autorreferidos", "Quantitativa contínua", "Insumo do IMC"],
        ["P034 / P035 / P04501", "Exercício e sedentarismo", "Qualitativa / quantitativa", "Preditora X"],
        ["P00901 / P015 / P018", "Verduras, peixe, frutas (dias/sem.)", "Quantitativa discreta", "Preditora X (anti-inflamatório)"],
        ["P01101 / P02501 / P02002", "Carne vermelha, doces, refrigerante", "Quantitativa discreta", "Preditora X (pró-inflamatório)"],
        ["N001", "Autoavaliação de saúde", "Qualitativa ordinal", "Preditora X / controle"],
        ["Q092", "Diagnóstico de depressão", "Qualitativa nominal", "Controle / preditora X"],
        ["I00101", "Possui plano de saúde", "Qualitativa nominal", "Controle / viés"],
    ], "Tabela 1 — Atributos selecionados da PNS 2019 a partir do Modelo Conceitual "
       "(extrato; ver classificação completa no Guia de Análises).",
       col_widths=[3.0, 5.0, 4.0, 4.0], font_size=7.5)
    write_here(doc)

    # 3.5 Etapa 4 — Missing
    h2(doc, "3.5 Etapa 4 — Tratamento de Dados Ausentes e Vazios")
    box(doc, "📝 O QUE ESCREVER AQUI — este é um dos seus pontos fortes", [
        ("txt", "A maioria dos trabalhos trata missing de forma genérica. Você fez algo mais "
                "sofisticado: distinguiu missing ESTRUTURAL (skip patterns do questionário) "
                "de missing ALEATÓRIO real. Reporte isso com destaque — é um diferencial "
                "metodológico."),
        ("sub", "Perguntas norteadoras:"),
        ("q", "O que é um 'skip pattern' do questionário da PNS? Dê o exemplo mais didático "
              "(quem responde 'não pratico exercício' não é perguntado 'quantos dias por "
              "semana' — o campo fica vazio sem ser dado faltante)."),
        ("q", "Por que tratar esse vazio como missing aleatório e imputá-lo pela média seria "
              "um ERRO metodológico?"),
        ("q", "Quantas variáveis foram excluídas por terem >75% de dados ausentes reais? "
              "Quais eram, e por que isso é uma limitação? (Renda, autoavaliação, "
              "escolaridade — perda da dimensão socioeconômica.)"),
        ("q", "Por que a imputação foi por média/moda DA PRÓPRIA CLASSE? Qual o risco disso "
              "para o ML? (Vazamento — refazer dentro dos folds de validação cruzada.)"),
    ])
    sk = load_csv(RES_COMOR / "tabelas" / "etapa2_5_skip_patterns.csv")
    add_df(doc, sk, "Tabela 3.4 — Skip patterns resolvidos: valores ausentes estruturais "
                    "preenchidos antes da auditoria de missing (desenho com comorbidades).",
           font_size=7.5)
    add_figure(doc, RES_COMOR / "figuras" / "etapa3_heatmap_missing.png",
               "Figura 3.1 — Mapa de calor de dados ausentes reais por variável e grupo, "
               "após a resolução dos skip patterns.", largura_cm=11.5)
    write_here(doc)

    # 3.6 Etapa 5 — Outliers
    h2(doc, "3.6 Etapa 5 — Análise e Remoção de Outliers")
    box(doc, "📝 O QUE ESCREVER AQUI", [
        ("sub", "Perguntas norteadoras:"),
        ("q", "Qual método de detecção você usou? (IQR — intervalo interquartil.)"),
        ("q", "Você usou 3×IQR em vez do 1,5×IQR padrão. Por quê? (Mais conservador — só "
              "remove valores extremos prováveis de erro de medição, não variação biológica.)"),
        ("q", "Por que o IQR foi calculado SEPARADAMENTE por classe (artrite vs. saudável)? "
              "O que aconteceria se fosse misturando os grupos?"),
        ("q", "Quantos outliers foram substituídos e em quais variáveis se concentraram?"),
    ])
    box(doc, "⚠️ ATENÇÃO — divergência com o template", [
        ("txt", "O texto-padrão do modelo do Prof. Zárate menciona '1,5×IQR'. Você usou "
                "3×IQR. NÃO copie o texto-padrão cego — ajuste para 3×IQR e justifique. "
                "Inconsistência entre o que você fez e o que escreveu é exatamente o que o "
                "avaliador procura."),
    ], cor=C_CUIDADO)
    out = load_csv(RES_COMOR / "tabelas" / "etapa4_log_outliers.csv")
    add_df(doc, out, "Tabela 3.5 — Log de detecção de outliers pelo método IQR×3,0 por classe "
                     "(desenho com comorbidades).", font_size=7.5)
    add_figure(doc, RES_COMOR / "figuras" / "etapa4_boxplots_outliers.png",
               "Figura 3.2 — Boxplots das variáveis-chave antes e depois da remoção de outliers.")
    write_here(doc)

    # 3.7 Etapa 6 — Preparação
    h2(doc, "3.7 Etapa 6 — Preparação dos Dados (Fusão, Transformação e Discretização)")
    box(doc, "📝 O QUE ESCREVER AQUI", [
        ("sub", "Gabarito: item 17 (fusão e discretização) + item 18 (dataset final). É a "
                "Tabela 2 do artigo."),
        ("sub", "Perguntas norteadoras:"),
        ("q", "Como o IMC foi criado e por que combinar peso e altura num único atributo é "
              "melhor do que mantê-los separados?"),
        ("q", "O que são os escores Inflamatório e Saudável? Que hipótese clínica eles "
              "operacionalizam?"),
        ("q", "Por que idade e IMC foram discretizados por limiares FIXOS (OMS) enquanto o "
              "escore inflamatório foi discretizado por QUARTIS? O que justifica métodos "
              "diferentes?"),
        ("q", "Qual o N e o número de atributos do dataset final? Você vai reportar UM "
              "desenho ou os DOIS (pura e com comorbidades)?"),
    ])
    add_rows_table(doc, ["Atributo gerado", "Origem / Regra de transformação", "Categorias resultantes"], [
        ["IMC", "Peso (kg) / Altura (m)² — fusão de P00104 e P00404", "Variável contínua"],
        ["Escore Inflamatório", "Soma de P01101, P02501, P02002, P02001 (hábitos pró-inflamatórios)", "Variável contínua"],
        ["Escore Saudável", "Soma de P018, P00901, P015, P023, P01601 (hábitos protetores)", "Variável contínua"],
        ["Razao_Inf_Saud", "Escore Inflamatório / (Escore Saudável + 1)", "Variável contínua"],
        ["FxEtaria_cat", "pd.cut em C008 por limiares OMS", "0=60–69 / 1=70–79 / 2=80+"],
        ["IMC_cat", "pd.cut em IMC por limiares OMS", "0=Baixo peso / 1=Normal / 2=Sobrepeso / 3=Obesidade"],
        ["AtivFisica_cat", "pd.cut em P035 por frequência", "0=Sedentário / 1=Pouco / 2=Moderado / 3=Ativo"],
        ["EscInfla_cat", "pd.qcut em Escore Inflamatório (quartis)", "0=Q1 / 1=Q2 / 2=Q3 / 3=Q4"],
    ], "Tabela 2 — Atributos finais após fusão, transformação e discretização "
       "(resultado dos notebooks 03 e 03B).", col_widths=[3.0, 7.5, 5.5], font_size=7.5)
    e9p = rel_puro["rastreabilidade"]["etapa_9_dataset_final"]
    e9c = rel_comor["rastreabilidade"]["etapa_9_dataset_final"]
    add_rows_table(doc, ["Desenho", "N registros", "N features", "Distribuição (Saudável / Artrite)", "Razão"], [
        ["Artrite pura", f"{e9p['n_registros']:,}".replace(",", "."), str(e9p["n_features"]),
         f"{e9p['distribuicao_label']['0']:,} / {e9p['distribuicao_label']['1']:,}".replace(",", "."),
         f"{e9p['razao_desbalanceamento']}:1"],
        ["Artrite com comorbidades", f"{e9c['n_registros']:,}".replace(",", "."), str(e9c["n_features"]),
         f"{e9c['distribuicao_label']['0']:,} / {e9c['distribuicao_label']['1']:,}".replace(",", "."),
         f"{e9c['razao_desbalanceamento']}:1"],
    ], "Tabela 3.6 — Conjuntos de dados finais dos dois desenhos de estudo.",
       col_widths=[3.8, 2.4, 2.2, 4.6, 3.0], font_size=8)
    write_here(doc)

    # 3.8 Etapa 7 — Seleção por entropia
    h2(doc, "3.8 Etapa 7 — Seleção de Atributos (Entropia)")
    box(doc, "⚠️ PENDÊNCIA — depende do notebook 04 (Machine Learning)", [
        ("txt", "A seleção de atributos por entropia de Shannon ainda não foi executada. "
                "Pertence ao notebook 04. Deixe esta seção marcada e volte a ela depois."),
        ("q", "Para já ir pensando: por que atributos com baixa entropia em relação à classe "
              "são pouco discriminativos e podem ser descartados?"),
    ], cor=C_PENDENTE)
    write_here(doc)

    # 3.9 Etapa 8 — Balanceamento
    h2(doc, "3.9 Etapa 8 — Balanceamento do Conjunto de Dados")
    box(doc, "📝 O QUE ESCREVER AQUI — PARCIAL (você já sabe as razões)", [
        ("txt", "O RUS (Random Under Sampler) em si será aplicado no notebook 04, mas você já "
                "SABE os números de desbalanceamento e pode descrever a estratégia."),
        ("sub", "Perguntas norteadoras:"),
        ("q", "Por que o desbalanceamento do desenho 'artrite pura' (8,77:1) é um problema "
              "sério, e como o RUS o resolve? Qual o custo de usar RUS?"),
        ("q", "Por que o RUS é aplicado APENAS no conjunto de treino, nunca no teste?"),
        ("q", "O desenho 'com comorbidades' (1,08:1) praticamente não precisa de "
              "balanceamento — como você reporta essa diferença entre os dois desenhos?"),
    ])
    box(doc, "⚠️ PENDÊNCIA — Tabela 3 (distribuição treino/teste pós-RUS)", [
        ("txt", "A Tabela 3 do artigo (divisão treino/teste após balanceamento) será "
                "preenchida com os números reais após o notebook 04."),
    ], cor=C_PENDENTE)
    para(doc, "[TABELA 3 — Inserir após notebook 04: divisão treino/teste pós-RUS]",
         italic=True, color=CINZA, justify=False)
    write_here(doc)

    # 3.10 e 3.11
    h2(doc, "3.10 Etapa 9 — Modelos de Aprendizado e Parametrização")
    box(doc, "⚠️ PENDÊNCIA — notebook 04", [
        ("txt", "Material de base: o Plano de Artigo já define os 4 algoritmos e a estratégia "
                "de validação."),
        ("q", "Por que essa mistura de 'caixa-branca' (Árvore de Decisão, Naive Bayes) e "
              "'caixa-preta' (Random Forest, AdaBoost)? Cite Loyola-González (2019)."),
        ("q", "Como funciona a otimização por RandomizedSearchCV?"),
    ], cor=C_PENDENTE)
    write_here(doc)
    h2(doc, "3.11 Etapa 10 — Treinamento e Teste")
    box(doc, "⚠️ PENDÊNCIA — notebook 04", [
        ("q", "Como será a divisão treino/teste e a validação cruzada 10-fold? Quais "
              "métricas (Precisão, Recall, F1, AUC-ROC)?"),
    ], cor=C_PENDENTE)
    write_here(doc)
    pagebreak(doc)

    # ===== 4. RESULTADOS =====
    h1(doc, "4. Experimentos e Análise dos Resultados")
    box(doc, "📝 DECISÃO ESTRUTURAL IMPORTANTE", [
        ("txt", "O template do Prof. Zárate espera que a Seção 4 contenha SÓ resultados de "
                "Machine Learning. Mas você produziu uma análise exploratória e bivariada "
                "rica que não tem 'slot' óbvio."),
        ("txt", "Recomendação: a sua análise descritiva/bivariada entra como subseção inicial "
                "— '4.1 Caracterização da Amostra' — antes dos resultados de ML. É o "
                "equivalente à 'Tabela 1' de artigos epidemiológicos. Discuta essa "
                "organização com o orientador, mas tê-la fortalece muito o artigo."),
        ("sub", "REGRA DE OURO da seção Resultados:"),
        ("txt", "Resultados NÃO interpreta — só reporta os números e o que as tabelas/figuras "
                "mostram. A interpretação vai para a Discussão. E nunca repita no texto o que "
                "já está na tabela."),
    ])

    h2(doc, "4.1 Caracterização da Amostra (análise descritiva e bivariada)")
    box(doc, "📝 O QUE ESCREVER AQUI — TUDO PRONTO (notebook 02)", [
        ("sub", "Perguntas norteadoras (lembre: só reportar, não interpretar):"),
        ("q", "Qual a prevalência global de artrite? Como varia por sexo, idade e escolaridade?"),
        ("q", "Quais variáveis diferiram significativamente entre artrite e saudáveis? Liste "
              "com seus p-valores (foram 6: sexo, autoavaliação de saúde, plano de saúde, "
              "IMC, frutas, refrigerante)."),
        ("q", "Quantas comorbidades, em média, tem um idoso com artrite? Quais as 3 mais "
              "prevalentes? (Hipertensão 65,3%, colesterol 39,8%, diabetes 21,9%.)"),
    ])
    tab3b = load_csv(RES_EDA / "tabelas" / "tabela_3B_prevalencia_subgrupos.csv")
    add_df(doc, tab3b, "Tabela 4.1 — Prevalência de artrite por subgrupo, com IC95% "
                       "(notebook 02).", font_size=7.5)
    add_figure(doc, RES_EDA / "figuras" / "grafico_3B_prevalencia_subgrupos.png",
               "Figura 4.1 — Prevalência de artrite por subgrupo com IC95%. Linha tracejada = "
               "prevalência global de 17,7%.")
    tab2c = load_csv(RES_EDA / "tabelas" / "tabela_2C_bivariada_completa.csv")
    add_df(doc, tab2c, "Tabela 4.2 — Análise bivariada completa: artrite pura vs. saudáveis, "
                       "com teste estatístico, p-valor e IC95% (notebook 02).", font_size=6.5)
    add_figure(doc, RES_EDA / "figuras" / "grafico_2A_boxplots_quantitativas.png",
               "Figura 4.2 — Boxplots das variáveis quantitativas por grupo (artrite pura vs. "
               "saudáveis).")
    add_figure(doc, RES_EDA / "figuras" / "grafico_2B_barras_qualitativas.png",
               "Figura 4.3 — Distribuição percentual das variáveis qualitativas por grupo.")
    comorb = load_csv(RES_COMOR / "tabelas" / "etapa3_5_comorbidades_prevalencia.csv")
    add_df(doc, comorb, "Tabela 4.3 — Prevalência de cada comorbidade entre os 4.025 idosos "
                        "com artrite (notebook 03B).", font_size=8,
           col_widths=[2.5, 5.0, 2.0, 2.0])
    add_figure(doc, RES_COMOR / "figuras" / "etapa3_5_comorbidades.png",
               "Figura 4.4 — Distribuição do número de comorbidades por idoso com artrite "
               "(esquerda) e prevalência de cada doença associada (direita).")
    write_here(doc)

    h2(doc, "4.2 Resultados do Treinamento dos Modelos")
    box(doc, "⚠️ PENDÊNCIA — notebook 04 (Machine Learning)", [
        ("txt", "Tabela 4 (F1-Score com IC95% no treinamento) — preencher após notebook 04."),
        ("q", "Para depois: qual algoritmo teve o melhor F1-Score na validação cruzada?"),
    ], cor=C_PENDENTE)
    para(doc, "[TABELA 4 — Inserir após notebook 04: F1-Measure com IC 95% por classe]",
         italic=True, color=CINZA, justify=False)
    write_here(doc)

    h2(doc, "4.3 Resultados do Conjunto de Teste")
    box(doc, "⚠️ PENDÊNCIA — notebook 04", [
        ("txt", "Gráfico 1 (barras de Precisão/Recall/F1) e Tabela 5 (métricas no teste) — "
                "após notebook 04."),
    ], cor=C_PENDENTE)
    para(doc, "[GRÁFICO 1 + TABELA 5 — Inserir após notebook 04]", italic=True, color=CINZA,
         justify=False)
    write_here(doc)

    h2(doc, "4.4 Análise do Modelo Final e Conhecimento Extraído")
    box(doc, "⚠️ PENDÊNCIA — notebook 04", [
        ("txt", "Gráfico 2 (Feature Importance) e as regras da Árvore de Decisão — após "
                "notebook 04."),
        ("q", "Para depois: os atributos de maior importância confirmam ou contrariam a "
              "hipótese nutricional inicial? Que regras a árvore extraiu?"),
    ], cor=C_PENDENTE)
    para(doc, "[GRÁFICO 2 + REGRAS — Inserir após notebook 04]", italic=True, color=CINZA,
         justify=False)
    write_here(doc)
    pagebreak(doc)

    # ===== 5. CONCLUSÃO =====
    h1(doc, "5. Conclusão e Trabalhos Futuros")
    box(doc, "📝 O QUE ESCREVER AQUI — Conclusão/Discussão (escrever em 4º lugar)", [
        ("sub", "Gabarito: item 27 (observações), item 28 (limitações), item 29 (trabalhos "
                "futuros)."),
        ("sub", "Material que você já tem — 5 discussões propostas (ver "
                "Resultados_Consolidados):"),
        ("b", "O paradoxo nutricional e a causalidade reversa (artríticos comem mais frutas — "
              "contraintuitivo)."),
        ("b", "O IMC como elo entre artrite e comorbidades."),
        ("b", "O perfil sociodemográfico: a artrite tem gênero e idade."),
        ("b", "O dilema metodológico dos dois desenhos de estudo."),
        ("b", "A multimorbidade como achado clínico central."),
        ("sub", "Perguntas norteadoras:"),
        ("q", "Em 2–3 frases: qual o achado principal do trabalho?"),
        ("q", "O objetivo declarado na Introdução e a conclusão dizem a MESMA coisa?"),
        ("q", "Quais as limitações honestas? (Design transversal, dados autorreferidos, "
              "vazamento de delineamento no desenho B, imputação por classe, "
              "desbalanceamento, variáveis descartadas por missing.)"),
        ("q", "Que trabalhos futuros fazem sentido? (Dados longitudinais, biomarcadores "
              "laboratoriais, replicar com PNS 2024.)"),
    ])
    box(doc, "⚠️ ATENÇÃO", [
        ("txt", "A Discussão não pode passar de 1/3 do artigo (critério explícito do Prof. "
                "Zárate). Se estiver longa, provavelmente você está repetindo resultados ou "
                "interpretando demais."),
    ], cor=C_CUIDADO)
    write_here(doc)
    pagebreak(doc)

    # ===== AGRADECIMENTOS / REFERÊNCIAS =====
    h1(doc, "Agradecimentos")
    box(doc, "📝 O QUE ESCREVER AQUI", [
        ("txt", "Texto-padrão de agradecimento (CNPq, FIP/PUC Minas, orientador). Ver "
                "modelo no template original."),
    ])
    write_here(doc)
    h1(doc, "Referências")
    box(doc, "📝 O QUE ESCREVER AQUI", [
        ("txt", "Padrão Vancouver (numérico). Mínimo 15–20 referências. Obrigatórias: "
                "IBGE/PNS 2019, CAPTO (Zárate et al. 2023), Loyola-González 2019, e os 6+ "
                "trabalhos relacionados."),
        ("b", "O Plano de Artigo (Tabela 26) e o Guia de Análises trazem a maioria das "
              "referências já levantadas — Malta et al., Szwarcwald et al., Claro et al., "
              "Louzada et al., Cancella & Zárate, Melo et al., Silva & Zárate, Xue et al."),
    ])
    write_here(doc)
    pagebreak(doc)

    # ===== PAINEL DE CONTROLE =====
    h1(doc, "Painel de Controle — o que já dá para escrever vs. o que falta")
    box(doc, "Use este painel para priorizar suas sessões de escrita", [
        ("sub", "PODE ESCREVER AGORA (não depende de mais nada):"),
        ("b", "Seção 3 inteira até a Etapa 6 (Materiais e Métodos — Etapas 1 a 6)."),
        ("b", "Seção 4.1 (Caracterização da Amostra) — todas as tabelas e figuras já estão "
              "inseridas neste documento."),
        ("b", "Seção 2 (Trabalhos Relacionados) — material no Plano de Artigo."),
        ("b", "Limitações da Conclusão — 6 limitações já levantadas."),
        ("sub", "DEPENDE DO NOTEBOOK 04 (Machine Learning):"),
        ("b", "Seção 3, Etapas 7 a 10 (seleção por entropia, RUS, modelos, treino/teste)."),
        ("b", "Seções 4.2, 4.3 e 4.4 (resultados de ML, feature importance, regras)."),
        ("b", "Tabelas 3, 4, 5 e Gráficos 1, 2 do artigo."),
        ("sub", "PENDÊNCIA MANUAL (não sai de código):"),
        ("b", "Figura 1 — Modelo Conceitual CAPTO. Desenhe a partir da Tabela 3.3 deste "
              "documento."),
        ("sub", "Estimativa:"),
        ("txt", "Com o material atual você consegue redigir aproximadamente 55–60% do artigo. "
                "O caminho crítico para fechar os 40% restantes é executar o notebook 04."),
    ], cor=C_PRONTO)

    doc.add_paragraph()
    fim = doc.add_paragraph(); fim.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = fim.add_run("— Template de trabalho gerado a partir dos notebooks 02, 03 e 03B —\n"
                     "A prosa é sua. As caixas azuis dizem o que escrever; o material já está "
                     "posicionado.")
    rf.italic = True; rf.font.size = Pt(8.5); rf.font.color.rgb = CINZA

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(f"[OK] Template de trabalho gerado: {OUT_PATH}")
    print(f"     {len(doc.paragraphs)} paragrafos, {len(doc.tables)} tabelas, "
          f"{len(doc.inline_shapes)} figuras")


if __name__ == "__main__":
    build()
